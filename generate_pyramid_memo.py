#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PYRAMID MEMO GENERATOR — Minto-structured written document, Copilot-executable.

Produces a .docx memo from a completed pyramid (SCQA + governing thought +
key line + support detail + conclusion + ask). Three archetypes, switched in
META["archetype"]:
  "recommendation" — structured memo, standalone reader (Minto default)
  "dot_dash"       — dot-dash outline / lap visual (author present in the room)
  "status"         — information / project-update memo (status -> variances ->
                     decisions needed; author the key line in that pattern)

HOW TO USE (in Copilot, Claude, or any assistant with a code interpreter):
  1. Attach this file. Say: "Run this script and give me the .docx it creates."
     It ships with fictional demonstration content (Meridian Office Services)
     so the structure is visible before you put your own argument in.
  2. To build your own memo: the assistant follows the STAGED PROMPTS below and
     edits ONLY the DATA section. Everything under BUILD — DO NOT EDIT stays.
  3. Re-run after every edit. The SELF-CHECK prints PASS/FAIL per rule and the
     script refuses to write a file while any check fails. Fix the DATA, never
     the checks.

THE SHARED SCHEMA: the PYRAMID data below is identical to the DATA section of
generate_presentation_deck.py and generate_standalone_deck.py. Fill it once,
paste it into all three, and one argument drives the memo and both decks.
run_pyramid_analysis.py produces a hand-off section that maps onto these
fields when the argument comes out of a diagnostic. Generator-specific
registers sit at the END of each DATA section (this memo: EVIDENCE_REGISTER;
the decks: EXHIBIT_REGISTER + MAX_VINTAGE_MONTHS) — when pasting the shared
block between scripts, keep each script's own registers in place.

STAGED PROMPTS — the assistant follows these in order and STOPS at each gate.

  STAGE 0 — APPLY THE CLIENT'S BRAND (optional). All brand values live in
  the THEME dict at the top of the DATA section; the shipped default is the
  neutral TMTH palette. Two routes to a client brand: (i) paste the client's
  palette hexes and font names into THEME; (ii) attach any existing client
  .pptx and run extract_theme.py (same repository) — it prints a
  ready-to-paste THEME dict with a suggested role mapping. Client brand
  arrives at runtime; never commit a client THEME to the public repository.

  STAGE 1 — FRAME. Ask the user, one question at a time: who is the reader and
  what do they already believe; what decision is wanted; which of the four
  patterns is this (giving direction / seeking approval / explaining how /
  choosing); will the author be present? Choose the archetype from the answer
  (present + short = dot_dash; standalone = recommendation; progress check =
  status). Show the frame. GATE: get a yes before Stage 2.

  STAGE 2 — SCQA AND GOVERNING THOUGHT. Draft Situation (noncontroversial,
  nod-test), Complication (the tension), Question (exactly one), and the
  governing thought in three drafts: the WHAT, then + the SO-WHAT it produces,
  then compressed to 25 words or fewer, active voice, verb first. Propose a
  tone order (SCA standard / ASC direct / CSA concerned) with a reason.
  GATE: show all drafts, get a yes.

  STAGE 3 — BUILD THE PYRAMID. Ask "why is the governing thought true?" and
  draft 2-4 key-line arguments as full-sentence assertions. Name the ordering
  principle (time / structure / degree). MECE test: can one exhibit prove two
  arguments (merge them); if all are true MUST the reader accept the answer
  (else an argument is missing). Under each, 2-5 supports from evidence the
  user supplies — never invented. GATE: show the pyramid + the audit, get a yes.

  STAGE 4 — ATTACK. For each key-line argument, state the strongest objection
  the named reader could raise and how the evidence answers it. Concede and
  reframe where the objection is right. Revise the ask to carry its own
  decision gate. GATE: show the objection table, get a yes.

  STAGE 5 — EVIDENCE SWEEP (ADVERSARIAL). Attach the engagement's source
  materials. For EVERY claim that carries weight — each key-line point and
  each load-bearing support claim: (1) find the best supporting evidence
  and cite the source; (2) actively search for DISCONFIRMING evidence and
  record what was looked for AND what was found, even when nothing was
  found; (3) mark the status honestly: EVIDENCED / GAP_RISK_ACCEPTED (with
  a stated reason) / GAP_OPEN. Fill the EVIDENCE_REGISTER. The build
  REFUSES while any row is GAP_OPEN, and every key-line point must have a
  register row. Risk-accepted gaps render in the memo as an explicit
  "stated limits" block — never silently absorbed. GATE: show the register,
  get a yes.

  STAGE 6 — FILL AND RUN. Write the approved content into the DATA section:
  prose for every support (memo body), bullets (dot-dash / decks), sources,
  conclusion in Minto's three parts, one ask with gate + owner + date. Replace
  every square-bracketed placeholder — the build fails while any remain. Run
  the script; hand the user the .docx and the self-check report.

ENGAGEMENT ASSETS: a filled EVIDENCE_REGISTER is client work product. It
lives in the client tenant. Never commit a filled register to the public
repository — the shipped register holds fictional demonstration content only.

VERSION NOTES
  v1.1 — 20 Aug 2026: brand moved from code to data (THEME dict, neutral
         TMTH default; validation gate on hex/fonts). Adversarial evidence
         sweep added as Stage 5 with the EVIDENCE_REGISTER, a GAP_OPEN
         build gate, per-section evidence footnotes and a rendered
         "stated limits" block for risk-accepted gaps.
  v1.0 — 20 Aug 2026: first published version.

Requires: python-docx.  Output: pyramid-memo.docx
"""

# ============================================================================
# DATA — EDIT THIS SECTION ONLY
# All demonstration content is fictional (Meridian Office Services).
# ============================================================================

# --- THEME: all brand values live here, as data. Ship = neutral TMTH palette.
# To apply a client's brand: (i) paste their palette hexes + font names below,
# or (ii) attach any existing client .pptx and run extract_theme.py (same
# repository) — it prints a ready-to-paste THEME dict. Hex: 6 chars, no '#',
# no alpha. Never commit a client THEME to the public repository.
THEME = {
    "dominant":      "2F4858",  # lead brand colour — section heads, table header fills
    "accent":        "4E7C90",  # secondary brand colour — sub-headings
    "highlight":     "C9A227",  # the one-highlight colour (used by the deck generators)
    "card_fill":     "EEF3F6",  # panel / label-cell background tint
    "ink":           "1F2A31",  # near-black body ink
    "ink_soft":      "3C4C57",  # softer body ink (used by the deck generators)
    "muted":         "6B7A85",  # secondary text, sources, footers
    "chart_compare": "B9C6CD",  # non-highlighted chart elements (deck generators)
    "font_heading":  "Cambria",
    "font_body":     "Calibri",
}

META = {
    "title": "Win renewals on response time, not price",
    "subtitle": "Recommendation memo — SME contract churn",
    "author_line": "Strategy team",
    "date": "20 Aug 2026",
    "audience": "Managing director and the executive committee",
    "presence": "standalone",        # "standalone" or "live" (informational)
    "archetype": "recommendation",   # "recommendation" | "dot_dash" | "status"
    "footer_tag": "Demonstration content — every name and number is fictional",
}

SCQA = {
    "tone_order": "SCA",  # SCA standard | ASC direct | CSA concerned
    "situation": ("Meridian serves 1,200 SME service contracts across the "
                  "South East, worth £5.1M of annual revenue. Renewal pricing "
                  "has been flat for two years."),
    "complication": ("Churn has risen from 22% to 35% in three years — a "
                     "£1.8M annual revenue impact. The lost-deal evidence "
                     "points away from price: objections fell 30% while churn "
                     "climbed."),
    "question": ("How does Meridian win renewals in a market that now judges "
                 "providers on response time?"),
}

GOVERNING_THOUGHT = ("Reposition Meridian around a response-time guarantee — "
                     "the factor deciding renewals — recovering churn to 25% "
                     "and £1.2M annual revenue within four quarters.")

KEY_LINE_ORDER = "degree"  # "time" | "structure" | "degree"

KEY_LINE = [
    {
        "tag": "DIAGNOSIS",
        "assertion": "Churn is perception-driven, not price-driven",
        "why_it_matters": ("Every remedy on the table prices the wrong "
                           "problem. The evidence separates what customers "
                           "say from what they do — and it clears price as "
                           "the cause."),
        "supports": [
            {
                "assertion": ("Current customers renew at stable rates while "
                              "competitor-bid churn climbs"),
                "prose": ("Existing-contract churn has held between 21% and "
                          "23% for eight quarters. A price problem hits "
                          "paying customers first, and it has not. The rise "
                          "is concentrated in renewals where a competitor "
                          "bid is on the table."),
                "bullets": [
                    "Existing-contract churn steady at 21–23% for eight quarters",
                    "Price pain reaches current payers first — it has not",
                    "The rise sits in renewals facing a competitor bid",
                ],
                "source": "Contract system renewal records, Q3 2024 to Q2 2026, n=1,180",
                "exhibit": None,
                "notes": ("Land the distinction between the two churn "
                          "populations before anyone mentions price. The "
                          "whole argument turns on it."),
            },
            {
                "assertion": ("Price objections fell 30% while churn rose 13 "
                              "points"),
                "prose": ("Recorded price objections fell from 46 to 32 a "
                          "quarter across three years. If price drove the "
                          "losses, objections would rise with churn. They "
                          "moved the opposite way."),
                "bullets": [
                    "Objections fell from 46 to 32 a quarter",
                    "Churn rose 13 points across the same period",
                    "The two series move in opposite directions",
                ],
                "source": "CRM objection log, 2023 to 2026",
                "exhibit": {"exhibit_id": "EX-OBJECTIONS",
                            "kind": "column",
                            "unit": "Price objections per quarter",
                            "categories": ["2023", "2024", "2025", "2026"],
                            "values": [46, 41, 37, 32],
                            "highlight": 3,
                            "source": "CRM objection log, 2023 to 2026"},
                "notes": ("If price were the cause, this line would rise "
                          "with churn. Let the divergence do the work — no "
                          "editorialising needed."),
            },
            {
                "assertion": ("All 10 lost Q4 deals cite the competitor "
                              "guarantee and none cite price"),
                "prose": ("Every lost-deal review from the last quarter "
                          "names the competitor's four-hour response "
                          "guarantee. Price appears in none of them. The "
                          "market has changed what it judges providers on."),
                "bullets": [
                    "10 of 10 reviews cite the response guarantee",
                    "Zero reviews cite price",
                    "Scope and relationship reasons trail far behind",
                ],
                "source": "Lost-deal reviews, Oct to Dec 2025, n=10",
                "exhibit": {"exhibit_id": "EX-LOSTDEAL",
                            "kind": "bar",
                            "unit": "Mentions across 10 lost-deal reviews",
                            "categories": ["Competitor response guarantee",
                                           "Service scope",
                                           "Relationship change",
                                           "Price"],
                            "values": [10, 3, 2, 0],
                            "highlight": 0,
                            "source": "Lost-deal reviews, Oct to Dec 2025, n=10"},
                "notes": ("This is the exhibit the sceptics remember. Ten "
                          "out of ten is unusual — say so, and say the n is "
                          "small but the signal is unanimous."),
            },
        ],
        "transition": ("If price is not the lever, the lost-deal records say "
                       "what is — and it is buildable."),
    },
    {
        "tag": "REMEDY",
        "assertion": ("A four-hour response guarantee wins the frame the "
                      "market now uses"),
        "why_it_matters": ("The competitor changed the question renewals are "
                           "judged on. Matching the guarantee answers that "
                           "question directly — and Meridian starts closer "
                           "to it than the sales narrative admits."),
        "supports": [
            {
                "assertion": ("The competitor guarantee reset what buyers ask "
                              "at renewal"),
                "prose": ("Renewal conversations now open with response "
                          "commitments, not price. Account managers report "
                          "the guarantee is raised by the customer in most "
                          "competitive renewals. The frame has moved and the "
                          "proposition has not moved with it."),
                "bullets": [
                    "Renewal conversations now open with response commitments",
                    "Customers raise the guarantee unprompted",
                    "The proposition has not moved with the frame",
                ],
                "source": "Account team renewal debriefs, Q1 to Q2 2026",
                "exhibit": None,
                "notes": ("Anchor this in the account managers' own words — "
                          "the evidence is qualitative and honest about it."),
            },
            {
                "assertion": ("Meridian already meets four hours on 78% of "
                              "callouts"),
                "prose": ("Dispatch records show 78% of 2025 callouts were "
                          "answered inside four hours without any guarantee "
                          "in place. The gap to close is the remaining 22% — "
                          "evenings, weekends and multi-site jobs. That gap "
                          "is known and priceable."),
                "bullets": [
                    "78% of 2025 callouts already inside four hours",
                    "The gap is evenings, weekends and multi-site jobs",
                    "The gap is known and priceable",
                ],
                "source": "Dispatch system records, Jan to Dec 2025, n=8,400",
                "exhibit": {"exhibit_id": "EX-CALLOUTS",
                            "kind": "bar",
                            "unit": "% of 2025 callouts",
                            "categories": ["Met within four hours today",
                                           "Requires new cover"],
                            "values": [78, 22],
                            "highlight": 0,
                            "source": "Dispatch system records, 2025, n=8,400"},
                "notes": ("The surprise is how close we already are. The "
                          "guarantee prices the last 22%, not the whole "
                          "operation."),
            },
            {
                "assertion": ("The guarantee costs £350k a year to stand up"),
                "prose": ("Three cost lines close the 22% gap: extended "
                          "engineering cover, a triage desk, and a service-"
                          "credit reserve priced at a 2% breach rate. The "
                          "total is £350k a year, fixed."),
                "bullets": [
                    "Extended engineering cover — £190k",
                    "Triage desk and dispatch tooling — £95k",
                    "Service-credit reserve at 2% breach — £65k",
                ],
                "source": "Operations costing model v2, Jul 2026",
                "exhibit": {"exhibit_id": "EX-COSTS",
                            "kind": "table",
                            "headers": ["Cost line", "Annual cost"],
                            "rows": [
                                ["Evening and weekend engineering cover", "£190k"],
                                ["Triage desk and dispatch tooling", "£95k"],
                                ["Service-credit reserve at 2% breach rate", "£65k"],
                                ["Total", "£350k"],
                            ],
                            "source": "Operations costing model v2, Jul 2026"},
                "notes": ("The reserve line is the honest one — a guarantee "
                          "without a priced breach rate is a slogan."),
            },
        ],
        "transition": ("A guarantee Meridian can meet at a known cost must "
                       "still pay. It does — on pessimistic assumptions."),
    },
    {
        "tag": "ECONOMICS",
        "assertion": ("The programme pays back within four quarters on the "
                      "central case"),
        "why_it_matters": ("The case is priced in ranges, not a point value. "
                           "Even the pessimistic case clears the outlay "
                           "inside five quarters, and the staged spend caps "
                           "the downside."),
        "supports": [
            {
                "assertion": ("Recovered renewals return £0.9M to £1.6M a "
                              "year against the £350k outlay"),
                "prose": ("The scenario table runs churn recovery at three "
                          "levels. The central case recovers churn to 27% "
                          "and returns £1.2M a year — payback inside four "
                          "quarters. The pessimistic case still pays back "
                          "in five."),
                "bullets": [
                    "Central case: churn to 27%, £1.2M a year recovered",
                    "Pessimistic case still pays back in five quarters",
                    "All three cases clear the £350k outlay",
                ],
                "source": "Renewal economics model, scenario run Jul 2026",
                "exhibit": {"exhibit_id": "EX-SCENARIOS",
                            "kind": "table",
                            "headers": ["Scenario", "Churn recovered to",
                                        "Revenue recovered", "Payback"],
                            "rows": [
                                ["Pessimistic", "31%", "£0.9M a year", "5 quarters"],
                                ["Central", "27%", "£1.2M a year", "4 quarters"],
                                ["Optimistic", "25%", "£1.6M a year", "3 quarters"],
                            ],
                            "source": "Renewal economics model, Jul 2026"},
                "notes": ("Never present the central case alone. The range "
                          "is the credibility."),
            },
            {
                "assertion": ("The 90-day gate caps the downside at £120k"),
                "prose": ("Spend is staged: £120k reaches the field before "
                          "the 90-day review. If lost-deal citations of the "
                          "competitor guarantee have not fallen below five a "
                          "quarter by then, the programme stops before the "
                          "remaining spend commits."),
                "bullets": [
                    "Only £120k committed before the 90-day review",
                    "Stop trigger: citations still above five a quarter",
                    "The remaining spend never commits on a failed signal",
                ],
                "source": "Programme spend schedule, draft Jul 2026",
                "exhibit": None,
                "notes": ("The gate is what makes this approvable — the "
                          "decision is reversible at a known cost."),
            },
        ],
        "transition": None,
    },
]

CONCLUSION = {
    "restate": ("The renewal problem is a framing problem, and the guarantee "
                "answers it — churn back to 25% and £1.2M of revenue "
                "recovered within four quarters."),
    "key_assumption": ("This rests on the market's response-time frame "
                       "persisting. If lost-deal citations shift away from "
                       "it, we review at the 90-day gate."),
    "next_action": ("The operations director stands up the build team on "
                    "1 Oct 2026."),
}

ASK = {
    "decision": ("Approve the response-time guarantee programme at £350k a "
                 "year"),
    "gate": ("Stop at 90 days if lost-deal citations of the competitor "
             "guarantee have not fallen below five a quarter"),
    "owner": "Operations director",
    "date": "1 Oct 2026",
}

NEXT_STEPS = [
    ("Confirm the service-credit terms with legal", "Commercial lead", "12 Sep 2026"),
    ("Brief the account teams on the renewal script", "Sales director", "19 Sep 2026"),
]

# --- EVIDENCE SWEEP (Stage 5) — one row per key-line point (claim_id = its
# tag) and per load-bearing support claim (claim_id = TAG-S<n>). status:
# "EVIDENCED" | "GAP_RISK_ACCEPTED" (status_reason required) | "GAP_OPEN"
# (blocks the build). counter_evidence_search records what was looked for
# AND what was found — even when nothing was found. Every key-line tag MUST
# have a row. Risk-accepted gaps render as the memo's "stated limits" block.
# ENGAGEMENT ASSET: never commit a filled register to the public repository.
EVIDENCE_REGISTER = [
    {"claim_id": "DIAGNOSIS",
     "claim": "Churn is perception-driven, not price-driven",
     "best_supporting_evidence": ("Existing-contract churn stable at 21–23% "
                                  "while competitor-bid churn climbs; price "
                                  "objections fell 30% as churn rose"),
     "source": ("Contract system renewal records n=1,180; CRM objection log "
                "2023–2026"),
     "counter_evidence_search": ("Looked for: price-led losses and discount "
                                 "pressure in the 2025–26 renewal notes. "
                                 "Found: two competitor rate cards below "
                                 "ours — and zero lost deals citing them."),
     "status": "EVIDENCED", "status_reason": ""},
    {"claim_id": "DIAGNOSIS-S3",
     "claim": ("All 10 lost Q4 deals cite the competitor guarantee and none "
               "cite price"),
     "best_supporting_evidence": ("Unanimous citation across the full Q4 "
                                  "2025 lost-deal review set"),
     "source": "Lost-deal reviews, Oct to Dec 2025, n=10",
     "counter_evidence_search": ("Looked for: selection bias in which lost "
                                 "deals received reviews. Found: none — all "
                                 "ten Q4 losses were reviewed, none "
                                 "skipped."),
     "status": "EVIDENCED", "status_reason": ""},
    {"claim_id": "REMEDY",
     "claim": ("A four-hour response guarantee wins the frame the market "
               "now uses"),
     "best_supporting_evidence": ("78% of 2025 callouts already inside four "
                                  "hours (n=8,400); the residual gap is "
                                  "costed at £350k a year"),
     "source": ("Dispatch system records 2025; operations costing model v2, "
                "Jul 2026"),
     "counter_evidence_search": ("Looked for: evidence the 22% gap is not "
                                 "closable at the quoted cost — external "
                                 "cover quotes for evenings and weekends. "
                                 "Found: quotes within the model's range."),
     "status": "EVIDENCED", "status_reason": ""},
    {"claim_id": "ECONOMICS",
     "claim": ("The programme pays back within four quarters on the central "
               "case"),
     "best_supporting_evidence": ("Scenario table: all three cases clear "
                                  "the £350k outlay; pessimistic case pays "
                                  "back in five quarters"),
     "source": "Renewal economics model, scenario run Jul 2026",
     "counter_evidence_search": ("Looked for: observed churn-recovery rates "
                                 "from comparable guarantee launches. "
                                 "Found: none in-house; market analogues "
                                 "are directional only."),
     "status": "GAP_RISK_ACCEPTED",
     "status_reason": ("Recovery rates are modelled, not observed — "
                       "untestable before launch. Bounded by the "
                       "pessimistic case and the 90-day stop gate.")},
]

# ============================================================================
# BUILD — DO NOT EDIT BELOW THIS LINE
# ============================================================================
import re
import sys

OUTPUT_FILE = "pyramid-memo.docx"

# ---------------------------------------------------------------- self-check
BANNED_LABELS = {"background", "analysis", "overview", "introduction",
                 "context", "summary", "next steps", "recommendation",
                 "findings", "update", "situation", "options", "approach",
                 "details", "conclusion"}

def _assertion_defect(text):
    """Return a defect string if the text reads as a label, else None."""
    t = (text or "").strip()
    words = t.split()
    if len(words) < 5:
        return "under 5 words — reads as a label, not a claim"
    if t.rstrip(".").lower() in BANNED_LABELS:
        return "banned label heading"
    if t.endswith(":"):
        return "ends in a colon — a label, not a claim"
    tail = [w for w in words[1:] if w[:1].isalpha()]
    caps = sum(1 for w in tail if w[:1].isupper())
    if tail and caps > 0.6 * len(tail):
        return "Title Case — use sentence case"
    return None

def _collect_strings(obj, out):
    if isinstance(obj, str):
        out.append(obj)
    elif isinstance(obj, dict):
        for v in obj.values():
            _collect_strings(v, out)
    elif isinstance(obj, (list, tuple)):
        for v in obj:
            _collect_strings(v, out)

def _long_sentences(strings, limit=25):
    hits = []
    for s in strings:
        for sent in re.split(r"(?<=[.!?])\s+", s):
            if len(sent.split()) > limit:
                hits.append(sent.strip()[:70] + "…")
    return hits

def _valid_exhibit(ex):
    if ex is None:
        return None
    if not isinstance(ex, dict):
        return "exhibit is not a dict"
    kind = ex.get("kind")
    if kind in ("bar", "column", "line"):
        cats, vals = ex.get("categories"), ex.get("values")
        if not cats or not vals or len(cats) != len(vals):
            return "chart categories/values missing or mismatched"
        hl = ex.get("highlight")
        if hl is not None and not (0 <= hl < len(vals)):
            return "highlight index out of range"
    elif kind == "table":
        if not ex.get("headers") or not ex.get("rows"):
            return "table needs headers and rows"
    else:
        return "exhibit kind must be bar / column / line / table"
    if not ex.get("source"):
        return "exhibit has no source"
    return None

SWEEP_STATUSES = ("EVIDENCED", "GAP_RISK_ACCEPTED", "GAP_OPEN")

def _theme_defects():
    defects = []
    for k, v in THEME.items():
        if k.startswith("font_"):
            if not str(v).strip():
                defects.append("THEME['%s'] is empty — name a real font" % k)
        elif not re.fullmatch(r"[0-9A-Fa-f]{6}", str(v)):
            defects.append("THEME['%s'] = %r is not a 6-char hex "
                           "(no '#', no alpha)" % (k, v))
    return defects

def run_self_check():
    checks, warnings = [], []

    td = _theme_defects()
    checks.append(("THEME: colours are 6-char hex and fonts are named"
                   + ("" if not td else " — " + "; ".join(td)), not td))

    gt_words = len(GOVERNING_THOUGHT.split())
    checks.append(("governing thought is 25 words or fewer (%d)" % gt_words,
                   0 < gt_words <= 25))
    checks.append(("governing thought is a claim, not a question",
                   "?" not in GOVERNING_THOUGHT))
    checks.append(("SCQA complete: situation, complication, question, answer",
                   all([SCQA.get("situation"), SCQA.get("complication"),
                        SCQA.get("question"), GOVERNING_THOUGHT])))
    low_sit = " " + SCQA["situation"].lower()
    checks.append(("situation carries no 'but' / 'however' (tension belongs "
                   "in the complication)",
                   " but " not in low_sit and "however" not in low_sit))
    checks.append(("exactly one question, ending in a question mark",
                   SCQA["question"].count("?") == 1
                   and SCQA["question"].strip().endswith("?")))
    checks.append(("tone order is SCA / ASC / CSA",
                   SCQA.get("tone_order") in ("SCA", "ASC", "CSA")))
    checks.append(("archetype is recommendation / dot_dash / status",
                   META.get("archetype") in ("recommendation", "dot_dash",
                                             "status")))
    checks.append(("key line holds 2–4 arguments (%d)" % len(KEY_LINE),
                   2 <= len(KEY_LINE) <= 4))
    checks.append(("ordering principle named: time / structure / degree",
                   KEY_LINE_ORDER in ("time", "structure", "degree")))

    label_defects = []
    for kl in KEY_LINE:
        d = _assertion_defect(kl.get("assertion", ""))
        if d:
            label_defects.append("key line '%s…': %s"
                                 % (kl.get("assertion", "")[:40], d))
        for sp in kl.get("supports", []):
            d = _assertion_defect(sp.get("assertion", ""))
            if d:
                label_defects.append("support '%s…': %s"
                                     % (sp.get("assertion", "")[:40], d))
    checks.append(("every heading is a full-sentence claim, not a label"
                   + ("" if not label_defects else
                      " — " + "; ".join(label_defects)),
                   not label_defects))

    checks.append(("every key-line argument carries 2–5 supports",
                   all(2 <= len(kl.get("supports", [])) <= 5
                       for kl in KEY_LINE)))
    kl_assertions = [kl["assertion"].strip().lower() for kl in KEY_LINE]
    checks.append(("no two key-line arguments make the same claim (MECE "
                   "overlap)", len(set(kl_assertions)) == len(kl_assertions)))
    restates = [kl for kl in KEY_LINE
                if kl["assertion"].strip().lower() in
                [s["assertion"].strip().lower() for s in kl["supports"]]]
    checks.append(("no key-line argument merely restates a support (summary, "
                   "not restatement)", not restates))
    checks.append(("every support cites a source",
                   all(sp.get("source") for kl in KEY_LINE
                       for sp in kl["supports"])))
    checks.append(("every support carries body prose",
                   all((sp.get("prose") or "").strip() for kl in KEY_LINE
                       for sp in kl["supports"])))
    if META.get("archetype") == "dot_dash":
        checks.append(("dot-dash mode: every support carries 1–3 bullets",
                       all(1 <= len(sp.get("bullets") or []) <= 3
                           for kl in KEY_LINE for sp in kl["supports"])))

    ex_defects = []
    for kl in KEY_LINE:
        for sp in kl["supports"]:
            d = _valid_exhibit(sp.get("exhibit"))
            if d:
                ex_defects.append("'%s…': %s" % (sp["assertion"][:35], d))
    checks.append(("every exhibit is well-formed with a source"
                   + ("" if not ex_defects else " — " + "; ".join(ex_defects)),
                   not ex_defects))

    checks.append(("conclusion has Minto's three parts: restate, key "
                   "assumption, next action",
                   all([CONCLUSION.get("restate"),
                        CONCLUSION.get("key_assumption"),
                        CONCLUSION.get("next_action")])))
    checks.append(("single ask with decision, gate, owner and date",
                   all([ASK.get("decision"), ASK.get("gate"),
                        ASK.get("owner"), ASK.get("date")])))

    # --- evidence sweep (Stage 5): adversarial register, gated
    row_defects = []
    for r in EVIDENCE_REGISTER:
        rid = r.get("claim_id", "?")
        for fld in ("claim", "best_supporting_evidence", "source",
                    "counter_evidence_search"):
            if not str(r.get(fld, "")).strip():
                row_defects.append("%s missing %s" % (rid, fld))
        if r.get("status") not in SWEEP_STATUSES:
            row_defects.append("%s status must be EVIDENCED / "
                               "GAP_RISK_ACCEPTED / GAP_OPEN" % rid)
        if (r.get("status") == "GAP_RISK_ACCEPTED"
                and not str(r.get("status_reason", "")).strip()):
            row_defects.append("%s is GAP_RISK_ACCEPTED with no stated "
                               "reason" % rid)
    checks.append(("every register row is complete: claim, best evidence, "
                   "source, counter-evidence search, valid status"
                   + ("" if not row_defects else
                      " — " + "; ".join(row_defects)), not row_defects))
    open_gaps = [r["claim_id"] for r in EVIDENCE_REGISTER
                 if r.get("status") == "GAP_OPEN"]
    checks.append(("no GAP_OPEN rows — every claim is EVIDENCED or "
                   "explicitly GAP_RISK_ACCEPTED with a reason"
                   + ("" if not open_gaps else
                      " — resolve or risk-accept: " + ", ".join(open_gaps)),
                   not open_gaps))
    reg_ids = set(r.get("claim_id") for r in EVIDENCE_REGISTER)
    reg_claims = set(str(r.get("claim", "")).strip().lower()
                     for r in EVIDENCE_REGISTER)
    missing_kl = [kl["tag"] for kl in KEY_LINE
                  if kl["tag"] not in reg_ids
                  and kl["assertion"].strip().lower() not in reg_claims]
    checks.append(("every key-line point has a register row (claim_id = "
                   "tag, or claim = exact assertion text)"
                   + ("" if not missing_kl else
                      " — missing rows: " + ", ".join(missing_kl)),
                   not missing_kl))

    strings = []
    for blob in (META, SCQA, GOVERNING_THOUGHT, KEY_LINE, CONCLUSION, ASK,
                 NEXT_STEPS, EVIDENCE_REGISTER):
        _collect_strings(blob, strings)
    placeholders = sorted(set(m for s in strings
                              for m in re.findall(r"\[[^\]]+\]", s)))
    checks.append(("no bracketed placeholders remain"
                   + ("" if not placeholders else
                      " — replace: " + ", ".join(placeholders)),
                   not placeholders))

    warnings.extend("sentence over 25 words: " + h
                    for h in _long_sentences(strings)[:6])
    for w in ("failing", "weak", "poor", "terrible", "underperforming"):
        if w in SCQA["situation"].lower():
            warnings.append("judgment word '%s' in the situation — the "
                            "situation must pass the nod test" % w)
    return checks, warnings

def print_report(checks, warnings, wrote):
    print("=" * 64)
    print("SELF-CHECK — %s" % OUTPUT_FILE)
    for name, ok in checks:
        print("  [%s] %s" % ("PASS" if ok else "FAIL", name))
    for w in warnings:
        print("  [WARN] %s" % w)
    print("=" * 64)
    if not all(ok for _, ok in checks):
        raise SystemExit("SELF-CHECK FAILED — no file written. Fix the DATA "
                         "section, not the checks.")
    if wrote:
        print("written", OUTPUT_FILE)

_checks, _warnings = run_self_check()
if not all(ok for _, ok in _checks):
    print_report(_checks, _warnings, wrote=False)
    sys.exit(1)

# ---------------------------------------------------------------- docx build
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def _rgb(h):
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

SLATE = _rgb(THEME["dominant"])
BLUE = _rgb(THEME["accent"])
GOLD = _rgb(THEME["highlight"])
INK = _rgb(THEME["ink"])
MUTE = _rgb(THEME["muted"])
TINT_HEX = THEME["card_fill"]
DOMINANT_HEX = THEME["dominant"]
HEAD_FONT, BODY_FONT = THEME["font_heading"], THEME["font_body"]

doc = Document()
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Inches(0.9)
    sec.left_margin = sec.right_margin = Inches(1.0)
normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(11)
normal.font.color.rgb = INK

def para(text="", size=11, color=INK, bold=False, italic=False,
         font=BODY_FONT, space_after=6, align=None, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.bold = bold
    r.font.italic = italic
    return p

def shade(cell, hex_fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)

def cell_text(cell, text, size=10.5, color=INK, bold=False, italic=False,
              font=BODY_FONT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.bold = bold
    r.font.italic = italic

def panel(rows):
    """Single-column shaded panel: list of (label, text)."""
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    t.autofit = False
    t.columns[0].width = Inches(1.5)
    t.columns[1].width = Inches(5.0)
    for i, (lab, text) in enumerate(rows):
        shade(t.cell(i, 0), TINT_HEX)
        shade(t.cell(i, 1), "FFFFFF")
        cell_text(t.cell(i, 0), lab, size=10, color=SLATE, bold=True,
                  font=HEAD_FONT)
        cell_text(t.cell(i, 1), text, size=10.5)
    return t

def data_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        shade(t.cell(0, j), DOMINANT_HEX)
        cell_text(t.cell(0, j), h, size=10, color=RGBColor(0xFF, 0xFF, 0xFF),
                  bold=True, font=HEAD_FONT)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell_text(t.cell(1 + i, j), str(v), size=10)
    return t

def _register_row(tag, assertion):
    for r in EVIDENCE_REGISTER:
        if (r.get("claim_id") == tag
                or str(r.get("claim", "")).strip().lower()
                == assertion.strip().lower()):
            return r
    return None

_ACCEPTED_GAPS = [r for r in EVIDENCE_REGISTER
                  if r["status"] == "GAP_RISK_ACCEPTED"]

def stated_limits_block():
    """Risk-accepted gaps render explicitly — never silently absorbed."""
    if not _ACCEPTED_GAPS:
        return
    para("Stated limits", size=12, color=SLATE, bold=True, font=HEAD_FONT,
         space_after=3)
    para("These gaps are accepted, with reasons — they are limits of the "
         "argument, not omissions from it.", size=9.5, color=MUTE,
         italic=True, space_after=4)
    panel([(r["claim_id"], "%s — %s" % (r["claim"], r["status_reason"]))
           for r in _ACCEPTED_GAPS])
    para("", space_after=8)

def exhibit_block(sp):
    ex = sp.get("exhibit")
    if not ex:
        return
    if ex["kind"] == "table":
        data_table(ex["headers"], ex["rows"])
    else:
        headers = [""] + [str(c) for c in ex["categories"]]
        rows = [[ex.get("unit", "value")] + [str(v) for v in ex["values"]]]
        data_table(headers, rows)
        para("Exhibit: build a %s chart from this table. Title = the claim "
             "above. Highlight: %s." % (
                 ex["kind"],
                 ex["categories"][ex["highlight"]]
                 if ex.get("highlight") is not None else "none"),
             size=9, color=MUTE, italic=True, space_after=2)
    para("Source: " + ex["source"], size=9, color=MUTE, italic=True,
         space_after=10)

# --- header
para(META["title"], size=20, color=SLATE, bold=True, font=HEAD_FONT,
     space_after=2)
para(META["subtitle"], size=12, color=BLUE, font=HEAD_FONT, space_after=2)
para("%s · %s · for: %s" % (META["author_line"], META["date"],
                            META["audience"]),
     size=9.5, color=MUTE, space_after=14)

# --- SCQA block (always first, per the authoring standard)
panel([("Situation", SCQA["situation"]),
       ("Complication", SCQA["complication"]),
       ("Question", SCQA["question"]),
       ("Answer", GOVERNING_THOUGHT)])
para("", space_after=4)

archetype = META["archetype"]

# --- introduction in the chosen tone order
def intro_text():
    s, c, a = SCQA["situation"], SCQA["complication"], GOVERNING_THOUGHT
    order = {"SCA": (s, c, a), "ASC": (a, s, c), "CSA": (c, s, a)}
    parts = order[SCQA["tone_order"]]
    return " ".join(parts)

if archetype == "dot_dash":
    para("Read together in the room — the outline carries the claims; the "
         "connective tissue is spoken.", size=9.5, color=MUTE, italic=True)
    para(" ".join([SCQA["situation"], SCQA["complication"]]), size=11,
         space_after=10)
    para(GOVERNING_THOUGHT, size=13, color=SLATE, bold=True, font=HEAD_FONT,
         space_after=10)
    for i, kl in enumerate(KEY_LINE):
        para("%d.  %s" % (i + 1, kl["assertion"]), size=12, color=SLATE,
             bold=True, font=HEAD_FONT, space_after=4)
        for sp in kl["supports"]:
            p = para("–  " + sp["assertion"], size=11, space_after=2)
            p.paragraph_format.left_indent = Inches(0.35)
            for b in sp.get("bullets", []):
                p2 = para("·  " + b, size=10.5, color=INK, space_after=1)
                p2.paragraph_format.left_indent = Inches(0.7)
            p3 = para("Evidence: " + sp["source"], size=9, color=MUTE,
                      italic=True, space_after=6)
            p3.paragraph_format.left_indent = Inches(0.7)
        reg = _register_row(kl["tag"], kl["assertion"])
        if reg:
            p4 = para("Register %s (%s): counter-evidence — %s"
                      % (reg["claim_id"], reg["status"],
                         reg["counter_evidence_search"]), size=8.5,
                      color=MUTE, italic=True, space_after=8)
            p4.paragraph_format.left_indent = Inches(0.35)
    para("", space_after=6)
    stated_limits_block()
    para("The ask", size=12, color=SLATE, bold=True, font=HEAD_FONT,
         space_after=4)
    panel([("Decision", ASK["decision"]), ("Gate", ASK["gate"]),
           ("Owner", ASK["owner"]), ("Date", ASK["date"])])
else:
    if archetype == "status":
        para("Status update — the sections run status against plan, "
             "variances, then the decisions needed.", size=9.5, color=MUTE,
             italic=True)
    para("Introduction", size=13, color=SLATE, bold=True, font=HEAD_FONT,
         space_after=4)
    para(intro_text(), size=11, space_after=8)
    preview = "; ".join("%d) %s" % (i + 1, kl["assertion"].rstrip("."))
                        for i, kl in enumerate(KEY_LINE))
    para("The argument, in %s order: %s." % (KEY_LINE_ORDER, preview),
         size=11, italic=True, space_after=12)

    for i, kl in enumerate(KEY_LINE):
        para("%d. %s" % (i + 1, kl["assertion"]), size=14, color=SLATE,
             bold=True, font=HEAD_FONT, space_after=4)
        para(kl["why_it_matters"], size=11, space_after=8)
        for sp in kl["supports"]:
            para(sp["assertion"], size=11.5, color=BLUE, bold=True,
                 font=HEAD_FONT, space_after=3)
            para(sp["prose"], size=11, space_after=4)
            exhibit_block(sp)
            if not sp.get("exhibit"):
                para("Evidence: " + sp["source"], size=9, color=MUTE,
                     italic=True, space_after=10)
        reg = _register_row(kl["tag"], kl["assertion"])
        if reg:
            para("Evidence line — register %s (%s): %s (%s). "
                 "Counter-evidence search: %s"
                 % (reg["claim_id"], reg["status"],
                    reg["best_supporting_evidence"], reg["source"],
                    reg["counter_evidence_search"]), size=8.5, color=MUTE,
                 italic=True, space_after=8)
        if kl.get("transition"):
            para(kl["transition"], size=11, italic=True, color=MUTE,
                 space_after=12)

    para("Conclusion", size=14, color=SLATE, bold=True, font=HEAD_FONT,
         space_after=4)
    para(CONCLUSION["restate"], size=11, space_after=4)
    para("Key assumption: " + CONCLUSION["key_assumption"], size=11,
         space_after=4)
    para("Next action: " + CONCLUSION["next_action"], size=11, space_after=12)

    stated_limits_block()

    para("The decision requested" if archetype != "status"
         else "Decisions needed", size=14, color=SLATE, bold=True,
         font=HEAD_FONT, space_after=4)
    panel([("Decision", ASK["decision"]), ("Gate", ASK["gate"]),
           ("Owner", ASK["owner"]), ("Date", ASK["date"])])
    if NEXT_STEPS:
        para("", space_after=2)
        para("Next steps (self-evident actions only)", size=11.5, color=SLATE,
             bold=True, font=HEAD_FONT, space_after=3)
        for n, (act, owner, date) in enumerate(NEXT_STEPS):
            para("%d. %s — %s, %s" % (n + 1, act, owner, date), size=10.5,
                 space_after=2)

para("", space_after=2)
para(META["footer_tag"], size=8.5, color=MUTE, italic=True,
     align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(OUTPUT_FILE)
print_report(_checks, _warnings, wrote=True)
