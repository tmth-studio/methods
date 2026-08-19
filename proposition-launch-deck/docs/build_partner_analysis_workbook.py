#!/usr/bin/env python3
"""Partner Decision Analysis Workbook — walks Copilot through the analyses
behind the partner-decision deck (archetype #2). Interview-driven; the final
stage fills the deck generator's DATA section and runs it."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import parse_xml

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

NAVY = RGBColor(0x00, 0x39, 0x5D)
CYAN = RGBColor(0x00, 0x83, 0xC5)
GREEN = RGBColor(0x2F, 0x6B, 0x4F)
RED = RGBColor(0xA8, 0x44, 0x3C)
MUTE = RGBColor(0x6E, 0x7F, 0x8D)

def shade(cell, hexfill):
    cell._tc.get_or_add_tcPr().append(parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{hexfill}"/>'))

def h1(t):
    h = doc.add_heading(t, level=1)
    for r in h.runs: r.font.color.rgb = NAVY
def h3(t, c=None):
    h = doc.add_heading(t, level=3)
    for r in h.runs: r.font.color.rgb = c or NAVY
def p(t, bold=False, italic=False, color=None, size=None):
    par = doc.add_paragraph()
    r = par.add_run(t); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
def bullets(items):
    for it in items: doc.add_paragraph(it, style='List Bullet')
def numbered(items):
    for it in items: doc.add_paragraph(it, style='List Number')

def box(label, lines, fill, lc):
    tbl = doc.add_table(rows=1, cols=1); tbl.style = 'Table Grid'
    c = tbl.cell(0, 0); shade(c, fill)
    r = c.paragraphs[0].add_run(label); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = lc
    for line in lines:
        rr = c.add_paragraph().add_run(line); rr.font.size = Pt(10)
    doc.add_paragraph()

def prompt(text):
    tbl = doc.add_table(rows=1, cols=1); tbl.style = 'Table Grid'
    c = tbl.cell(0, 0); shade(c, 'EFEBF7')
    r = c.paragraphs[0].add_run('PROMPT FOR COPILOT — copy the full block into Copilot chat')
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x5B, 0x21, 0xB6)
    rr = c.add_paragraph().add_run(text)
    rr.font.name = 'Consolas'; rr.font.size = Pt(8.5)
    rr._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    doc.add_paragraph()

def output_table(headers, rows=4):
    par = doc.add_paragraph()
    r = par.add_run('OUTPUT — record the approved result here')
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = CYAN
    tbl = doc.add_table(rows=rows+1, cols=len(headers)); tbl.style = 'Table Grid'
    for j, htxt in enumerate(headers):
        c = tbl.cell(0, j); shade(c, 'D9EEF9')
        rr = c.paragraphs[0].add_run(htxt); rr.bold = True; rr.font.size = Pt(9)
    for i in range(1, rows+1):
        for j in range(len(headers)):
            tbl.cell(i, j).paragraphs[0].add_run(' ').font.size = Pt(10)
    doc.add_paragraph()

def gate(t):
    tbl = doc.add_table(rows=1, cols=1); tbl.style = 'Table Grid'
    c = tbl.cell(0, 0); shade(c, 'D9EEF9')
    r = c.paragraphs[0].add_run('GATE: ' + t); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = NAVY
    doc.add_paragraph()

# ================= COVER =================
t = doc.add_heading('Partner Decision Analysis Workbook', level=0)
for r in t.runs: r.font.color.rgb = NAVY
p('The analysis behind the partner-decision deck, run with Copilot — stage by stage, interview first.', size=12)
p('Companion to generate_partner_decision_deck.py (archetype #2: "should we partner to launch this feature?"). The final stage fills the generator\'s DATA section and runs it — the script\'s self-check is the reconcile step.', italic=True, color=MUTE, size=9)

h1('How to use this workbook')
numbered([
    'Attach this document to a Copilot chat, together with generate_partner_decision_deck.py and (if demand is not yet scored) objective-customer-value-standard.md.',
    'Work one stage at a time, in order. Paste the stage prompt. Copilot interviews you first, then produces the output.',
    'Check the output against the stage checklist. Record it in the output table. It is the input for the next stage.',
    'Stage 5 writes the results into the generator\'s DATA section and runs it. The script computes the crossover, checks the recommendation against the numbers, and hands back the finished deck.',
])
box('RULES FOR EVERY STAGE — COPILOT MUST FOLLOW THESE', [
    '1. Interview first. If information is missing or unclear, ask numbered questions and STOP. Do not produce output around a gap.',
    '2. Never invent data. Every number carries a source and a date in the evidence register.',
    '3. Evidence for AND against. In this workbook that cuts both ways: the case FOR partnering and the case AGAINST both get built before either is judged.',
    '4. Confidence and delivery estimates come from the delivery team and the partner\'s committed dates — never from the strategy side.',
    '5. End every output with a self-check table: rule | PASS or FAIL | repair made.',
], 'FDE8E8', RED)

# ================= STAGE 0 =================
doc.add_page_break()
h1('Stage 0 — Setup: define the decision')
p('Copilot interviews you to establish exactly what is being decided, and what the capacity would otherwise do.', italic=True, color=MUTE)
h3('What Copilot asks you for', GREEN)
bullets([
    'The feature: name, one-line description, the segment it serves',
    'Its demand status: already scored (attach the backlog analysis) or not yet scored (run OCV first)',
    'The candidate routes: build, partner, buy/white-label — and any the client has ruled out, with reasons',
    'Candidate partners: names, or where a shortlist would come from',
    'The capacity context: what slot/team would this consume, and what is its NEXT BEST USE (this becomes the do-nothing floor)',
    'Hurdle rate, horizon, and any deal constraints (procurement rules, regulatory approvals, existing vendor exclusivities)',
])
prompt("""You run Stage 0 (Setup) of the partner decision analysis, using the attached Partner Decision Analysis Workbook.

Interview me. Ask these questions in order, ONE GROUP AT A TIME, and wait for my answers:
1. FEATURE: What feature is this decision about? One line on what it does and for whom.
2. DEMAND STATUS: Has demand been scored (OCV or equivalent)? If yes, paste the score, the gate and the evidence base. If no, we run the OCV standard first - say so and stop after this interview.
3. ROUTES: Which routes are on the table - build, partner, buy? Has anything been ruled out already, and on what evidence?
4. PARTNERS: Which partners are candidates? If none named yet, where would a shortlist come from?
5. THE FLOOR: If we do NOT launch this feature, what does the capacity do instead, and what is that worth? (This becomes the do-nothing floor - the number every route must beat.)
6. CONSTRAINTS: Hurdle rate, horizon, procurement or regulatory constraints, existing exclusivities.

Then produce the DECISION REGISTER: feature + segment | demand status and score | routes in scope with any exclusions | partner candidates | the floor and its basis | constraints | evidence inventory (item, source, date) | gaps as TO COLLECT with owners.

Rules: interview first; never invent data; the floor needs a stated basis, not a guess. End with the self-check table.""")
box('YOUR REVIEW CHECKLIST', [
    '✓ The floor is real: the next best use of the capacity is named and valued, not assumed to be zero.',
    '✓ Ruled-out routes have stated evidence — a route excluded by preference re-enters the analysis.',
    '✓ Demand status is honest: an unscored feature goes through OCV before this analysis continues.',
    '✓ Every constraint that could kill the deal (procurement, regulatory, exclusivity) is on the register now, not discovered at signing.',
], 'FFF8E6', RGBColor(0x8A, 0x6D, 0x0F))
output_table(['Item', 'Value / source', 'Gap? Owner + date'], rows=6)
gate('The decision register is complete. The floor has a basis. Demand is scored or queued for scoring.')

# ================= STAGE 1 =================
doc.add_page_break()
h1('Stage 1 — Worth launching now: demand and the window')
p('Confirms demand and — the stage that makes this a route decision — prices the cost of being late. Feeds deck slides 9–10.', italic=True, color=MUTE)
h3('Inputs', GREEN)
bullets(['Demand score and evidence (Stage 0 / backlog analysis / fresh OCV run)',
         'Customer timing data: renewal dates, contract cycles, seasonal buying moments — from CRM or research',
         'Uptake model: how adoption falls if launch arrives after the window'])
prompt("""You run Stage 1 (Demand and the window) of the partner decision analysis.

Decision register: [PASTE from Stage 0]
Timing data and uptake model: [PASTE / attach]

Do this:
1. Confirm the demand case in two sentences: score vs gate, evidence base, stated limits. Do not re-run the scoring - reference it.
2. Find the WINDOW: when does the segment actually make this decision? Renewal moments, contract cycles, budget cycles - from data, not assumption. State what share of the segment's decision moments fall in the next N quarters.
3. Price being late: model uptake if launch arrives inside the window vs after it. State the mechanism (customers locked into new contracts, competitor entrenchment) - not just the number.
4. State the null test: if there is NO meaningful window, say so plainly. Without a window, build-vs-partner is a margin preference and this deck's argument changes.

Output: the demand confirmation, the window (share of decision moments + dates), the late-arrival uptake penalty with its mechanism, and the self-check table.""")
box('YOUR REVIEW CHECKLIST', [
    '✓ The window comes from customer timing data (renewals, cycles) — not from an internal deadline dressed up as a market fact.',
    '✓ The late-arrival penalty has a mechanism, not just a number.',
    '✓ The null test was answered honestly: if there is no window, the deck\'s key line 1 changes and speed stops being worth margin.',
], 'FFF8E6', RGBColor(0x8A, 0x6D, 0x0F))
output_table(['Element', 'Result', 'Source'], rows=4)
gate('Demand confirmed by reference. The window is evidenced and the cost of missing it is priced with a mechanism.')

# ================= STAGE 2 =================
doc.add_page_break()
h1('Stage 2 — Price every route, including doing nothing')
p('The economics core: all routes on one basis, risk-adjusted. Feeds deck slides 7 and 12.', italic=True, color=MUTE)
h3('Inputs', GREEN)
bullets(['Window and uptake penalty (Stage 1)', 'Delivery estimates per route, with ranges — from the delivery team',
         'Indicative partner terms (margin share, integration cost)', 'Unit economics: margin per customer per route', 'The floor (Stage 0)'])
prompt("""You run Stage 2 (Route pricing) of the partner decision analysis.

Decision register + window analysis: [PASTE from Stages 0-1]
Delivery estimates, partner terms, unit economics: [PASTE / attach]

For EVERY route in scope - build, partner, buy, AND do-nothing:
1. Price it: time to live (with range), capex, margin structure (full margin / revenue share / unit fee), and confidence from the delivery team or the partner's committed dates.
2. Apply the window: uptake inside vs after it (Stage 1's penalty). This is where the build route usually loses - show the mechanism, not just the discount.
3. Compute risk-adjusted NPV per route on ONE consistent basis. State the basis once and reuse it.
4. The do-nothing route gets the floor from Stage 0 as its value.
5. Decompose the winner's margin vs time trade explicitly: "route X gives up £[Y] per customer and still wins £[Z], because..." - the objection this pre-empts is 'we are giving away margin'.

Output: the route table (route | time | capex | margin structure | risk-adjusted NPV | verdict), the margin-vs-time decomposition, and the self-check table. The recommended route must be the highest NPV - if it is not, say which soft factor you are invoking and flag it for my decision.""")
box('YOUR REVIEW CHECKLIST', [
    '✓ One NPV basis across all routes — a route priced on different assumptions is not comparable.',
    '✓ Confidence bands trace to the delivery team and partner commitments, not to the strategy side.',
    '✓ The do-nothing floor is in the table and at least one route beats it.',
    '✓ The margin-vs-time decomposition is explicit — the "giving away margin" objection is answered before it is raised.',
    '✓ If the recommendation is not the highest NPV, the soft factor is named and it is YOUR call, not Copilot\'s.',
], 'FFF8E6', RGBColor(0x8A, 0x6D, 0x0F))
output_table(['Route', 'Time ±', 'Capex', 'Margin structure', 'Risk-adj NPV', 'Verdict'], rows=4)
gate('Every route priced on one basis. The floor is beaten. The winner wins on arithmetic — or the exception is escalated.')

# ================= STAGE 3 =================
doc.add_page_break()
h1('Stage 3 — Partner due diligence')
p('Names the partner and verifies them — including the negatives. Feeds deck slide 14.', italic=True, color=MUTE)
h3('Inputs', GREEN)
bullets(['Partner candidates (Stage 0)', 'Public filings, ratings, regulatory record', 'Reference calls with the partner\'s existing bank clients', 'The partner\'s committed integration timeline and terms'])
prompt("""You run Stage 3 (Partner due diligence) of the partner decision analysis.

Candidates: [PASTE from Stage 0]
Filings, references, terms: [PASTE / attach]

Do this:
1. Build the shortlist (2-3 candidates) with the selection criteria stated BEFORE scoring: integration speed, segment fit, terms flexibility, financial strength, conduct record.
2. For the leading candidate, verify: scale and track record | credit/product performance through a downturn | platform and integration references | conduct, complaints and regulatory record | capital/funding position.
3. Collect the NEGATIVES deliberately: what did references criticise? What is the concentration or funding risk? An assessment without stated negatives is not diligence.
4. Name the runner-up and the gap: why the leader over them, in one sentence per criterion.
5. Anything unverifiable at this stage becomes a condition precedent in the heads of terms - list them.

Output: the diligence card (who they are / what we verified / negatives / why them over the runner-up), the conditions-precedent list, and the self-check table.""")
box('YOUR REVIEW CHECKLIST', [
    '✓ Criteria were stated before scoring — not fitted to the preferred candidate.',
    '✓ The negatives are real and specific. "None found" means the diligence has not been done.',
    '✓ The runner-up comparison is honest enough that you could defend choosing them instead.',
    '✓ Unverified items became conditions precedent, not assumptions.',
], 'FFF8E6', RGBColor(0x8A, 0x6D, 0x0F))
output_table(['Element', 'Finding', 'Source / reference'], rows=5)
gate('The partner is named, verified, and has stated negatives. Unverifiables are conditions precedent.')

# ================= STAGE 4 =================
doc.add_page_break()
h1('Stage 4 — The crossover and the bounds')
p('Computes the re-internalise trigger and sets the four bounds. Feeds deck slides 13 and 16–17.', italic=True, color=MUTE)
h3('Inputs', GREEN)
bullets(['Margin per customer: in-house vs partnered (Stage 2)', 'Capex to build in-house later', 'Payback policy (years)', 'The deal terms under negotiation'])
prompt("""You run Stage 4 (Crossover and bounds) of the partner decision analysis.

Route economics: [PASTE from Stage 2]
Diligence card + conditions: [PASTE from Stage 3]
Build capex, margin figures, payback policy: [PASTE]

Do this:
1. Compute the crossover: build capex / (margin delta per customer x payback years) = the customer count at which building in-house pays back. Show the arithmetic. This is the re-internalise trigger.
2. Set the four bounds:
   a. What the partner holds (their capability, their risk)
   b. What we keep (customer, data, pricing, brand - the moat assets)
   c. Contract bounds (exclusivity scope, data ownership on exit, portability period, price re-opener)
   d. The re-internalise trigger from step 1, with its review cadence
3. Build the sensitivity table: partner slips / partner captures the relationship / margin share creeps / demand softer than scored. Each row: what changes, and the observable early signal.
4. Test the demand row specifically: does soft demand change the WINNER or only the size of the win? If it changes the winner, the recommendation is fragile - say so.

Output: the crossover arithmetic, the four bounds, the sensitivity table, the fragility verdict, and the self-check table.""")
box('YOUR REVIEW CHECKLIST', [
    '✓ The crossover is arithmetic you could repeat on paper — capex, delta, payback, result.',
    '✓ "What we keep" covers the moat assets: customer, data, pricing, brand. Missing one is how partners become competitors.',
    '✓ Every sensitivity row has an observable early signal.',
    '✓ The fragility test was run: you know whether soft demand changes the winner or only the margin of victory.',
], 'FFF8E6', RGBColor(0x8A, 0x6D, 0x0F))
output_table(['Element', 'Result'], rows=5)
gate('The trigger is computed. The four bounds are set. The recommendation is robust — or its fragility is stated.')

# ================= STAGE 5 =================
doc.add_page_break()
h1('Stage 5 — Generate the deck')
p('Fills the generator\'s DATA section and runs it. The script\'s self-check is the reconcile step.', italic=True, color=MUTE)
h3('The DATA map', GREEN)
tbl = doc.add_table(rows=8, cols=2); tbl.style = 'Table Grid'
rows = [('Workbook output', 'Generator DATA field'),
        ('Stage 0 — feature, segment, partner, floor', 'META · SCQ · ROUTES (do-nothing row)'),
        ('Stage 1 — window + late penalty', 'DEMAND_CARD · WAIT_STATS · SCQ complication'),
        ('Stage 2 — route table + decomposition', 'ROUTES · KEYLINE 2 · route economics rows'),
        ('Stage 3 — diligence card + conditions', 'DILIGENCE_ROWS · appendix items'),
        ('Stage 4 — crossover inputs + bounds + sensitivity', 'CROSSOVER · BOUNDS_ROWS · SENSITIVITY_ROWS'),
        ('The recommendation in one sentence (≤25 words)', 'GT · REC_LINES'),
        ('Owners and dates for the four asks', 'ASKS'),]
for i, (a, b) in enumerate(rows):
    ca, cb = tbl.cell(i, 0), tbl.cell(i, 1)
    if i == 0:
        shade(ca, 'D9EEF9'); shade(cb, 'D9EEF9')
    ra = ca.paragraphs[0].add_run(a); ra.font.size = Pt(9.5); ra.bold = (i == 0)
    rb = cb.paragraphs[0].add_run(b); rb.font.size = Pt(9.5); rb.bold = (i == 0)
doc.add_paragraph()
prompt("""You run Stage 5 (Generate the deck) of the partner decision analysis.

All approved stage outputs: [PASTE Stages 0-4 output tables]
Attached: generate_partner_decision_deck.py

Do this:
1. Edit ONLY the DATA section of the script. Map each approved output to its field using the workbook's DATA map. Write the governing thought as ONE sentence of 25 words or fewer: partner + time + NPV + the counterfactual + the bounds clause.
2. Do not touch the LAYOUT section. Do not change the self-check.
3. Run the script. Report its full self-check output to me verbatim - every PASS/FAIL line and the placeholder list.
4. If any check FAILS, fix the DATA (not the check) and re-run. A recommendation the numbers do not support must go back to Stage 2, not be forced through.
5. When all checks PASS: give me the generated .pptx, and list every remaining placeholder so I can see what is still dummy.

End with: the self-check output, the file, the placeholder list.""")
box('YOUR REVIEW CHECKLIST', [
    '✓ The script\'s self-check is all PASS — this replaces the manual reconcile checks entirely.',
    '✓ The placeholder list is empty, or every remaining bracket is deliberate.',
    '✓ Open the deck once: the route matrix (7), the crossover (13) and the bounds (16) read as your analysis, not the worked example.',
    '✓ The conditions precedent from Stage 3 made it into the terms and the asks.',
], 'FFF8E6', RGBColor(0x8A, 0x6D, 0x0F))
gate('Self-check all PASS. The deck is generated from your data, and the numbers cannot disagree with each other.')

out = '/tmp/methods/proposition-launch-deck/Partner_Analysis_Workbook.docx'
doc.save(out)
print('Saved:', out)
