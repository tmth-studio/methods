#!/usr/bin/env python3
"""Backlog Analysis Workbook — walks Copilot through the analyses behind
the proposition-backlog deck (archetype #6). Interview-driven: Copilot asks,
the user answers, outputs land in tables that feed named deck slides."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import parse_xml

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

SLATE = RGBColor(0x2F, 0x48, 0x58)
BLUE = RGBColor(0x4E, 0x7C, 0x90)
GOLDD = RGBColor(0x8A, 0x6D, 0x0F)
GREEN = RGBColor(0x2F, 0x6B, 0x4F)
RED = RGBColor(0xA8, 0x44, 0x3C)
MUTE = RGBColor(0x64, 0x81, 0x8F)
INK = RGBColor(0x1F, 0x2A, 0x31)

def shade(cell, hexfill):
    cell._tc.get_or_add_tcPr().append(parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{hexfill}"/>'))

def h1(t):
    h = doc.add_heading(t, level=1)
    for r in h.runs: r.font.color.rgb = SLATE
def h2(t):
    h = doc.add_heading(t, level=2)
    for r in h.runs: r.font.color.rgb = BLUE
def h3(t, c=None):
    h = doc.add_heading(t, level=3)
    for r in h.runs: r.font.color.rgb = c or INK
def p(t, bold=False, italic=False, color=None, size=None):
    par = doc.add_paragraph()
    r = par.add_run(t); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
def bullets(items):
    for it in items: doc.add_paragraph(it, style='List Bullet')
def numbered(items):
    for it in items: doc.add_paragraph(it, style='List Number')

def box(label, lines, fill, lc):
    tbl = doc.add_table(rows=1, cols=1); tbl.style = 'Table Grid'
    c = tbl.cell(0, 0); shade(c, fill)
    r = c.paragraphs[0].add_run(label); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = lc
    for line in lines:
        rr = c.add_paragraph().add_run(line); rr.font.size = Pt(10)
    doc.add_paragraph()

def prompt(text):
    tbl = doc.add_table(rows=1, cols=1); tbl.style = 'Table Grid'
    c = tbl.cell(0, 0); shade(c, 'EFEBF7')
    r = c.paragraphs[0].add_run('PROMPT FOR COPILOT — copy the full block into Copilot chat')
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x5B, 0x21, 0xB6)
    rr = c.add_paragraph().add_run(text)
    rr.font.name = 'Consolas'; rr.font.size = Pt(8.5)
    rr._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    doc.add_paragraph()

def output_table(headers, rows=4, label='OUTPUT — record the approved result here'):
    par = doc.add_paragraph()
    r = par.add_run(label); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = BLUE
    tbl = doc.add_table(rows=rows+1, cols=len(headers)); tbl.style = 'Table Grid'
    for j, htxt in enumerate(headers):
        c = tbl.cell(0, j); shade(c, 'DBE7F8')
        rr = c.paragraphs[0].add_run(htxt); rr.bold = True; rr.font.size = Pt(9)
    for i in range(1, rows+1):
        for j in range(len(headers)):
            tbl.cell(i, j).paragraphs[0].add_run(' ').font.size = Pt(10)
    doc.add_paragraph()

def gate(t):
    tbl = doc.add_table(rows=1, cols=1); tbl.style = 'Table Grid'
    c = tbl.cell(0, 0); shade(c, 'DBE7F8')
    r = c.paragraphs[0].add_run('GATE: ' + t); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = SLATE
    doc.add_paragraph()

# ================= COVER =================
t = doc.add_heading('Backlog Analysis Workbook', level=0)
for r in t.runs: r.font.color.rgb = SLATE
p('The analysis behind the proposition-backlog deck, run with Copilot — stage by stage, interview first.', size=12)
p('Companion to proposition-backlog-template.pptx (archetype #6: "what do we build, in what order?"). Each stage produces the numbers for named slides. The deck shows a worked example (exporting SMEs); your outputs replace it.', italic=True, color=MUTE, size=9)

h1('How to use this workbook')
numbered([
    'Attach this document to a Copilot chat, together with objective-customer-value-standard.md (same repository).',
    'Work one stage at a time, in order. Paste the stage prompt. Copilot interviews you first, then produces the output.',
    'Check the output against the stage checklist. Record the approved result in the output table. It is the input for the next stage.',
    'Stage 6 maps every number to its deck slide and runs the reconcile checks. Then use the deck’s own Copilot slide (slide 3) to insert the data.',
])
box('RULES FOR EVERY STAGE — COPILOT MUST FOLLOW THESE', [
    '1. Interview first. If information is missing or unclear, ask numbered questions and STOP. Do not produce output around a gap.',
    '2. Never invent data. Every number carries a source and a date in the evidence register. A number without a source does not enter the analysis.',
    '3. Evidence for AND against. A verdict reached from friendly evidence alone is not a verdict.',
    '4. End every output with a self-check table: rule | PASS or FAIL | repair made. Repair each FAIL before presenting the output.',
    '5. Say which deck slides the output feeds. The slide numbers are stated in each stage.',
], 'FDE8E8', RED)

# ================= STAGE 0 =================
doc.add_page_break()
h1('Stage 0 — Setup: define the engagement')
p('Copilot interviews you to establish what is being analysed. Nothing else can start before this.', italic=True, color=MUTE)
h3('What Copilot asks you for', GREEN)
bullets([
    'The segment: definition, size, revenue today, share of wallet',
    'The candidate features: name, one-line description, intended price point (rough is fine)',
    'Capacity: build slots per cycle, cycle length',
    'The earnings mandate and the hurdle rate',
    'Evidence available: customer research, campaign history, brand tracking, delivery estimates, partner conversations',
])
prompt("""You run Stage 0 (Setup) of the proposition backlog analysis, using the attached Backlog Analysis Workbook.

Interview me. Ask these questions in order, ONE GROUP AT A TIME, and wait for my answers:
1. SEGMENT: How do you define the segment? How many customers, what revenue today, what share of their wallet do you hold?
2. FEATURES: List every candidate feature: name, one line on what it does, intended price point. Include the ones you suspect are weak - the analysis decides, not the intake.
3. CAPACITY: How many build slots per cycle? How long is a cycle? What counts as one slot?
4. MONEY: What is the earnings mandate for this backlog? What hurdle rate applies?
5. EVIDENCE: What do you already hold - customer research (n, dates), past campaign results, brand tracking, delivery estimates, partner term sheets?

Then produce the ENGAGEMENT REGISTER:
- Segment definition and baseline numbers
- Feature list (each with description and price point)
- Capacity and cycle definition
- Mandate and hurdle
- Evidence inventory: item | source | date | covers which stage (1-5)
- Gaps: evidence we do not hold | which stage it blocks | who could get it

Rules: interview first; never invent data; flag every gap as TO COLLECT with an owner. End with the self-check table.""")
box('YOUR REVIEW CHECKLIST', [
    '✓ The segment definition is one you could defend to the client’s exec.',
    '✓ Every candidate feature is in - including the politically favoured weak ones. The analysis drops them, not the intake.',
    '✓ A "slot" has a concrete definition (a squad-cycle, a quarter of a team - something countable).',
    '✓ Every evidence item has a source and a date. Every gap has an owner.',
], 'FFF8E6', GOLDD)
output_table(['Item', 'Value / source', 'Gap? Owner + date'], rows=6)
gate('The engagement register is complete. Every gap has an owner. No analysis starts on an unowned gap.')

# ================= STAGE 1 =================
doc.add_page_break()
h1('Stage 1 — Demand: score and gate the features')
p('Produces the OCV demand scores and the admission verdicts. Feeds deck slides 8–9 and the matrix on slide 7.', italic=True, color=MUTE)
h3('Inputs from Stage 0', GREEN)
bullets(['Feature list with descriptions and price points', 'Segment research (interviews, n, dates)', 'Incumbent solution facts: what the segment uses today, at what cost', 'Past-launch calibration data (for the gate)'])
prompt("""You run Stage 1 (Demand) of the proposition backlog analysis. The OCV standard (attached) is the authority; follow it exactly.

Engagement register: [PASTE from Stage 0]
Research and incumbent facts: [PASTE / attach]

Work ONE FEATURE PER CHAT if the register is large. For each feature:
1. Derive the segment's outcomes: outcome ladder for the job, minimal covering set (OCV standard, step 1a).
2. Score the feature against the incumbent per component: financial, time and stress deltas, time priced at the standard's conventions.
3. Subtract switching costs (standard's taxonomy). Adjust each component for visibility and credibility - contractual anchors beat claimed ones.
4. Sum to the OCV score. Map to predicted uptake with the calibration table.
5. Set the gate from past launches - state the calibration, not a chosen number. Verdict per feature: ADMITTED or REFUSED.

Rules: interview first if any input is missing; evidence for AND against each score; every number into the evidence register with source and date. End with: the score table (feature | OCV | predicted uptake | vs gate | verdict), the gate and its calibration, and the self-check table.""")
box('YOUR REVIEW CHECKLIST', [
    '✓ The gate comes from calibration (past launches), not from the answer anyone wanted.',
    '✓ Each score has evidence for AND against. An empty against-column means Copilot did not look.',
    '✓ Features near the gate are flagged for re-test at every re-score.',
    '✓ Refused features have their reason stated plainly - they get slide 7 rows, not silence.',
], 'FFF8E6', GOLDD)
output_table(['Feature', 'OCV score', 'Predicted uptake', 'Vs gate', 'Verdict'], rows=6)
gate('Every feature scored with sourced evidence. The gate is calibrated. Admitted features go to Stages 2–4; refused features stop here.')

# ================= STAGE 2 =================
doc.add_page_break()
h1('Stage 2 — The wrapper: memory at the entry points')
p('Produces the wrapper case: the recall gap, the design, the break-even and the stop rule. Feeds deck slide 10 and one bridge bar on slide 13.', italic=True, color=MUTE)
h3('Inputs', GREEN)
bullets(['Entry-point research: when does the job arise, in the customer’s words', 'Brand tracking / recall data, yours and competitors’', 'Media and channel cost cards'])
prompt("""You run Stage 2 (Wrapper) of the proposition backlog analysis.

Engagement register: [PASTE]  Admitted features: [PASTE from Stage 1]
Entry-point research and tracking data: [PASTE / attach]

Do this:
1. Identify the segment's category entry points: the moments the job arises, phrased in the customer's words, from research - not from the marketing plan. If the research does not show them, STOP and say what study is needed.
2. Measure current unaided recall at each entry point, for the client and competitors. The gap is the wrapper's job description.
3. Propose the wrapper: one name, a fixed set of brand assets, explicit linkage to the entry points. Span ONLY the admitted features.
4. Price it: fixed annual cost against a recall-uplift band. State the break-even in recall points, with the arithmetic shown.
5. State the stop rule: the recall movement required by which month, or the spend stops.

Rules: interview first; a case that only asserts awareness value is refused - the break-even must be arithmetic. End with the wrapper card (gap / design / economics / stop rule) and the self-check table.""")
box('YOUR REVIEW CHECKLIST', [
    '✓ Entry points come from research, in customer words - not from the marketing plan.',
    '✓ The break-even is shown as arithmetic you could repeat on paper.',
    '✓ The stop rule names a number and a month.',
    '✓ The wrapper spans only admitted features - nothing below the gate gets wrapped.',
], 'FFF8E6', GOLDD)
output_table(['Element', 'Result'], rows=5)
gate('The wrapper case has a measured gap, a break-even in recall points, and a stop rule with a date.')

# ================= STAGE 3 =================
doc.add_page_break()
h1('Stage 3 — Configuration: price, incentives and defaults')
p('Produces the base case and the lever sizing per admitted feature — the earnings bridge. Feeds deck slides 13–14.', italic=True, color=MUTE)
h3('Inputs', GREEN)
bullets(['Predicted uptake per feature (Stage 1)', 'Margin model per feature', 'Past campaign results by lever type', 'Price elasticity evidence, channel cost cards'])
prompt("""You run Stage 3 (Configuration) of the proposition backlog analysis.

Admitted features + uptake: [PASTE from Stage 1]
Margin model, campaign history, elasticity evidence: [PASTE / attach]

For each admitted feature:
1. Base case: predicted uptake x price x margin, before any lever. Show the arithmetic.
2. Long-list levers across the 4Ps. Classify each: default / friction removal / price move / paid incentive / channel shift.
3. Two-part test on every PAID lever: (a) would the behaviour happen anyway? (b) could a default achieve it free? Estimate dead-weight from past campaigns. A lever fails if either answer is yes.
4. Size the survivors: uplift x margin - cost including dead-weight. NAME the rejected levers; their earnings stay out of the forecast.
5. Assemble the earnings bridge: base -> each lever add -> incentive cost back -> wrapper uplift (Stage 2) -> configured case. The bars must sum.

Rules: interview first; dead-weight estimates need a stated basis; the bridge arithmetic must reconcile exactly. End with the bridge table, the rejected-lever list with reasons, and the self-check table.""")
box('YOUR REVIEW CHECKLIST', [
    '✓ The bridge sums: base + adds − costs = configured case, exactly.',
    '✓ Every paid lever passed the two-part test; the dead-weight basis is stated.',
    '✓ Rejected levers are named with reasons - that is what makes the accepted ones credible.',
    '✓ Defaults are used wherever they work: free beats paid.',
], 'FFF8E6', GOLDD)
output_table(['Bridge element', '£m', 'Basis / source'], rows=7)
gate('The bridge reconciles. Every lever has passed or failed the two-part test on stated evidence.')

# ================= STAGE 4 =================
doc.add_page_break()
h1('Stage 4 — Routes: build, buy or partner, per feature')
p('Produces route pricing, the timing discounts, the counterfactual and the partner bounds. Feeds deck slides 15 and 18–19.', italic=True, color=MUTE)
h3('Inputs', GREEN)
bullets(['Capability register (or build it here)', 'Delivery team estimates with ranges — from the delivery team, not the strategy team', 'Partner candidates with indicative terms', 'Hurdle rate and horizon (Stage 0)'])
prompt("""You run Stage 4 (Routes) of the proposition backlog analysis.

Admitted features + configured cases: [PASTE from Stage 3]
Capability register, delivery estimates, partner terms: [PASTE / attach]

For each admitted feature:
1. Inventory the capabilities it requires - capabilities, not tasks. Keep one set of names end to end.
2. Classify each capability have / buy / build. Name real partner candidates with indicative terms - an unnamed partner is a build with optimism attached.
3. Price each viable route: time to launch, capex, margin share, confidence band. Confidence comes from the delivery team's and the partner's committed dates.
4. Discount the configured case for each route's timing and confidence. Keep the BUILT-INSTEAD counterfactual for any partnered feature - it prices the route decision.
5. For partner routes, set the bounds: what they hold, what we keep, contract terms (exclusivity, data, exit, re-opener), and the re-internalise trigger with its review point.
6. State each feature's slot consumption on its chosen route.

Rules: interview first; no unnamed partners; the counterfactual row is mandatory. End with the route table (feature | route | time | confidence | risk-adjusted NPV | slots), the partner bounds card, and the self-check table.""")
box('YOUR REVIEW CHECKLIST', [
    '✓ Confidence bands trace to the delivery team, not to hope.',
    '✓ Every partner is named with real indicative terms.',
    '✓ The built-instead counterfactual exists for every partnered feature.',
    '✓ Partner bounds cover: hold / keep / contract terms / re-internalise trigger.',
], 'FFF8E6', GOLDD)
output_table(['Feature', 'Route', 'Time ±', 'Risk-adj NPV', 'Slots'], rows=5)
gate('Every route priced with delivery-team confidence. Counterfactuals and partner bounds recorded.')

# ================= STAGE 5 =================
doc.add_page_break()
h1('Stage 5 — The sort: order the backlog and set the re-score')
p('Produces the sorted backlog, the cut-line, the governing-thought numbers and the re-score triggers. Feeds deck slides 5–7, 12 and 16.', italic=True, color=MUTE)
h3('Inputs', GREEN)
bullets(['Risk-adjusted NPV and slot consumption per feature (Stage 4)', 'Capacity per cycle (Stage 0)'])
prompt("""You run Stage 5 (Sort) of the proposition backlog analysis.

Route table: [PASTE from Stage 4]   Capacity: [PASTE from Stage 0]

Do this:
1. Compute the sort key per feature: risk-adjusted NPV / build slots consumed. Rank descending.
2. State where per-slot order differs from raw NPV order, and why - this is the slide-12 story.
3. Apply the cut-line: this cycle's slots fill from the top. Statuses: THIS CYCLE / NEXT CYCLE / OFF THE LIST (refused at Stage 1).
4. Compute the governing-thought numbers: total configured earnings of this cycle's slate, margin, by year N. These must equal the Stage 3 bridge outputs.
5. Set the re-score: cadence, and the triggers that can move the order (new demand evidence, return assumptions move, a capability lands, a slot frees). Each trigger: observable signal | who watches | mechanical consequence.
6. Draft the governing thought in ONE sentence, 25 words or fewer: the rule + this cycle's slate + the earnings so-what.

Rules: arithmetic shown; the GT sentence counted; triggers must be observable. End with the sorted backlog table, the GT sentence with word count, the trigger list, and the self-check table.""")
box('YOUR REVIEW CHECKLIST', [
    '✓ The per-slot arithmetic is shown and the ranking follows it.',
    '✓ The GT is one sentence, 25 words or fewer, and its earnings equal the bridge.',
    '✓ Every trigger has an observable signal and a named watcher.',
    '✓ The cut-line follows capacity - nothing jumped the queue.',
], 'FFF8E6', GOLDD)
output_table(['Position', 'Feature', 'Return per slot', 'Status'], rows=5)
gate('The backlog is sorted, cut and gated. The governing thought is one sentence that reconciles with the bridge.')

# ================= STAGE 6 =================
doc.add_page_break()
h1('Stage 6 — Handoff: fill the deck')
p('Maps every output to its slide, runs the reconcile checks, then hands to the deck’s own Copilot slide.', italic=True, color=MUTE)
h3('The slide map', GREEN)
tbl = doc.add_table(rows=9, cols=2); tbl.style = 'Table Grid'
rows = [('Workbook output', 'Deck slide'),
        ('Stage 0 — segment, capacity, mandate', 'Slides 4 (S-C-Q) and 1 (cover)'),
        ('Stage 1 — scores, gate, verdicts', 'Slides 8–9, matrix on 7'),
        ('Stage 2 — wrapper card', 'Slide 10, one bar on 13'),
        ('Stage 3 — bridge, lever verdicts', 'Slides 13–14'),
        ('Stage 4 — route table, bounds, counterfactual', 'Slides 15, 18–19'),
        ('Stage 5 — sorted backlog, GT, triggers', 'Slides 5–7, 12, 16, and the ask on 21'),
        ('Evidence register (all stages)', 'Sources lines throughout + appendix'),
        ('Gaps still open', 'Stated limits, slide 9'),]
for i, (a, b) in enumerate(rows):
    ca, cb = tbl.cell(i, 0), tbl.cell(i, 1)
    if i == 0:
        shade(ca, 'DBE7F8'); shade(cb, 'DBE7F8')
    ra = ca.paragraphs[0].add_run(a); ra.font.size = Pt(9.5); ra.bold = (i == 0)
    rb = cb.paragraphs[0].add_run(b); rb.font.size = Pt(9.5); rb.bold = (i == 0)
doc.add_paragraph()
prompt("""You run Stage 6 (Handoff) of the proposition backlog analysis.

All approved stage outputs: [PASTE Stages 0-5 output tables]

Do this:
1. Produce the slide-fill sheet: for each deck slide named in the workbook's slide map, list the exact values and text that replace the worked example.
2. Run the reconcile checks and report PASS or FAIL for each:
   a. Bridge total (slide 13) = governing thought earnings (slides 4-5)
   b. Return per slot (slide 12) = risk-adjusted NPV (slide 15) / slots
   c. Demand scores identical on slides 7 and 8
   d. Statuses on slide 7 match the sections and the ask on slide 21
3. List every number that still has no source in the evidence register.
4. STOP. Do not fill the deck from this chat. The deck is filled in PowerPoint using its own Copilot slide (slide 3), with this slide-fill sheet as the client data.

End with the slide-fill sheet, the reconcile report, and the self-check table.""")
box('YOUR REVIEW CHECKLIST', [
    '✓ Every reconcile check is PASS. An open FAIL never reaches PowerPoint.',
    '✓ Every number on the fill sheet has a register source.',
    '✓ Remaining gaps appear as stated limits, not as silence.',
], 'FFF8E6', GOLDD)
gate('Reconcile checks all PASS. Hand the slide-fill sheet to Copilot in PowerPoint with the deck’s slide 3.')

out = '/tmp/methods/proposition-launch-deck/Backlog_Analysis_Workbook.docx'
doc.save(out)
print('Saved:', out)
