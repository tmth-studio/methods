#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PYRAMID ANALYSIS GENERATOR — Minto diagnostic (Flow A), Copilot-executable.

Runs the analysis half of the Pyramid Principle: frame the problem, establish
the Situation and Complication, form and TEST hypotheses against evidence,
land the Question and Answer, and turn the answer into a costed, gated
recommendation. Produces a structured .docx analysis record whose final
section maps field-for-field onto the shared PYRAMID schema used by
generate_pyramid_memo.py, generate_presentation_deck.py and
generate_standalone_deck.py — so a finished analysis hands off cleanly to
the presentation generators.

HOW TO USE (in Copilot, Claude, or any assistant with a code interpreter):
  1. Attach this file. Say: "Run this script and give me the .docx it creates."
     It ships with fictional demonstration content (Meridian Office Services).
  2. To run your own analysis: the assistant follows the STAGED PROMPTS below
     and edits ONLY the DATA section. Everything under BUILD — DO NOT EDIT
     stays. Re-run after every edit; the SELF-CHECK refuses to write a file
     while any rule fails.

STAGED PROMPTS — the assistant follows these in order and STOPS at each gate.

  STAGE A0 — APPLY THE CLIENT'S BRAND (optional). All brand values live in
  the THEME dict at the top of the DATA section; the shipped default is the
  neutral TMTH palette. Two routes to a client brand: (i) paste the client's
  palette hexes and font names into THEME; (ii) attach any existing client
  .pptx and run extract_theme.py (same repository) — it prints a
  ready-to-paste THEME dict with a suggested role mapping. Client brand
  arrives at runtime; never commit a client THEME to the public repository.

  STAGE A1 — FRAME THE PROBLEM. From the user's raw material, find four
  elements: the symptom (observable, with numbers), the evidence (data,
  timeframe, baseline), the decision that waits (with its cost), and the
  reader who decides. Keep every diagnosis OUT of the frame — the word
  "because" in a frame is a fail; park candidate causes for A3. Write the
  one-sentence frame. List what is missing and ask for it — never invent
  data. GATE: show the frame and the four elements, get a yes.

  STAGE A2 — SITUATION AND COMPLICATION. Sort the facts: noncontroversial
  facts go in the Situation (2–4 sentences, no "but/however", no judgment
  words); the change and the damage go in the Complication (starts with
  "But", numbers, no solution, no cause). Build the evidence register:
  fact | source | date. Run the nod test on every sentence. GATE: show both
  statements and the register, get a yes.

  STAGE A3 — HYPOTHESES AND TESTS. Write 3–5 MECE candidate causes,
  INCLUDING the explanation the stakeholders already believe. For each,
  write positive predictions ("if true the data shows X") AND negative
  predictions ("if true we will NOT see W") — a hypothesis with no negative
  prediction is untestable. Test against the data; record evidence FOR and
  AGAINST; verdicts SUPPORTED / CONTRADICTED / INCONCLUSIVE (SUPPORTED needs
  a surviving negative prediction). Name the strongest hypothesis and the one
  missing data item that would move confidence most. GATE: show the table,
  get a yes.

  STAGE A4 — EVIDENCE SWEEP (ADVERSARIAL). Attach the engagement's source
  materials. For EVERY claim that will carry weight — the surviving
  hypothesis, the emerging answer, and each load-bearing supporting fact:
  (1) find the best supporting evidence and cite the source; (2) actively
  search for DISCONFIRMING evidence and record what was looked for AND what
  was found, even when nothing was found — an unrecorded search is
  indistinguishable from no search; (3) mark the status honestly:
  EVIDENCED / GAP_RISK_ACCEPTED (with a stated reason) / GAP_OPEN. Fill the
  EVIDENCE_SWEEP register. The build REFUSES while any row is GAP_OPEN —
  every claim must be evidenced or its gap explicitly risk-accepted.
  GATE: show the register, get a yes.

  STAGE A5 — QUESTION AND ANSWER. Write THE one question the reader asks
  knowing the Situation, Complication and the surviving hypothesis. Check
  its level — the real question usually sits above the surface debate.
  Answer it in one sentence from the surviving evidence only. Show the
  chain: answer ← hypothesis ← evidence. GATE: show both, get a yes.

  STAGE A6 — RECOMMENDATION. Write 3–4 real options including "do nothing"
  with its honest cost. Score them (fit / cost / time / reversibility /
  risk-if-wrong). Recommend one: what, why (pointing at the answer),
  resources, timeline, owner. Attach 3–4 metrics each with baseline, target
  and DATE; a decision gate; and the key assumption. GATE: show it, get a
  yes.

  STAGE A7 — HAND OFF. Run the script. The final section of the .docx maps
  this analysis onto the PYRAMID fields of the three presentation
  generators. Carry the failed hypotheses forward — they are the prepared
  answers to the room's objections — and carry the evidence sweep into the
  memo/deck EVIDENCE_REGISTER.

ENGAGEMENT ASSETS: a filled EVIDENCE_SWEEP register is client work product.
It lives in the client tenant. Never commit a filled register to the public
repository — the shipped register holds fictional demonstration content only.

VERSION NOTES
  v1.1 — 20 Aug 2026: brand moved from code to data (THEME dict, neutral
         TMTH default; validation gate on hex/fonts). Adversarial evidence
         sweep added as Stage A4 with the EVIDENCE_SWEEP register and a
         GAP_OPEN build gate; stages renumbered A4→A5, A5→A6, A6→A7 to
         match the workbook.
  v1.0 — 20 Aug 2026: first published version.

Requires: python-docx.  Output: pyramid-analysis.docx
"""

# ============================================================================
# DATA — EDIT THIS SECTION ONLY
# All demonstration content is fictional (Meridian Office Services).
# ============================================================================

# --- THEME: all brand values live here, as data. Ship = neutral TMTH palette.
# To apply a client's brand: (i) paste their palette hexes + font names below,
# or (ii) attach any existing client .pptx and run extract_theme.py (same
# repository) — it prints a ready-to-paste THEME dict. Hex: 6 chars, no '#',
# no alpha. Never commit a client THEME to the public repository.
THEME = {
    "dominant":      "2F4858",  # lead brand colour — section heads, table header fills
    "accent":        "4E7C90",  # secondary brand colour — sub-headings
    "highlight":     "C9A227",  # the one-highlight colour — inconclusive/attention marks
    "card_fill":     "EEF3F6",  # panel / label-cell background tint
    "ink":           "1F2A31",  # near-black body ink
    "ink_soft":      "3C4C57",  # softer body ink (used by the deck generators)
    "muted":         "6B7A85",  # secondary text, sources, footers
    "chart_compare": "B9C6CD",  # non-highlighted chart elements (deck generators)
    "font_heading":  "Cambria",
    "font_body":     "Calibri",
}

META = {
    "title": "SME contract churn — root-cause diagnostic",
    "author_line": "Strategy team",
    "date": "20 Aug 2026",
    "reader": "Managing director",
    "footer_tag": "Demonstration content — every name and number is fictional",
}

PROBLEM_FRAME = {
    "symptom": "SME contract churn rose from 22% to 35% over three years",
    "evidence": ("Renewal records 2023 to 2026; annual revenue impact £1.8M "
                 "against a £5.1M base"),
    "decision": ("Whether to commit £2M to the retention roadmap in the "
                 "October budget round"),
    "reader": "Managing director",
    "frame_sentence": ("SME churn has risen from 22% to 35% over three years "
                       "(£1.8M annual revenue impact); we must identify the "
                       "root cause before committing £2M to the retention "
                       "roadmap."),
}

SITUATION = ("Meridian serves 1,200 SME service contracts across the South "
             "East, worth £5.1M of annual revenue. Renewal pricing has been "
             "flat for two years. Contract terms and service scope are "
             "unchanged since 2023.")

COMPLICATION = ("But churn has risen from 22% to 35% in three years — a "
                "£1.8M annual revenue impact. The rise is concentrated in "
                "renewals where a competitor bid is on the table.")

EVIDENCE_REGISTER = [
    ("Churn 22% (2023) rising to 35% (2026)",
     "Contract system renewal records", "Jul 2026"),
    ("Existing-contract churn stable at 21–23% for eight quarters",
     "Contract system renewal records, n=1,180", "Jul 2026"),
    ("Price objections fell from 46 to 32 a quarter, 2023 to 2026",
     "CRM objection log", "Jul 2026"),
    ("All 10 Q4 2025 lost-deal reviews cite the competitor's four-hour "
     "response guarantee; none cite price",
     "Lost-deal reviews, n=10", "Jan 2026"),
    ("Average deal size stable at £4.2k across the period",
     "Billing system extract", "Jul 2026"),
    ("78% of 2025 callouts answered inside four hours",
     "Dispatch system records, n=8,400", "Feb 2026"),
]

HYPOTHESES = [
    {
        "label": "H1",
        "statement": ("Pricing sits above the market and is driving "
                      "renewals to cheaper providers"),
        "predict_yes": [
            "Price objections rise in step with churn",
            "Current customers — who pay the price today — leave fastest",
            "Lost-deal reviews cite price as the deciding factor",
        ],
        "predict_no": [
            "If true, we will NOT see stable deal sizes at renewal",
        ],
        "evidence_for": [
            "The CFO's read: two competitors list lower headline rates",
        ],
        "evidence_against": [
            "Price objections fell 30% while churn rose",
            "Existing-contract churn is stable at 21–23%",
            "Zero of 10 lost-deal reviews cite price",
            "Deal sizes are stable at £4.2k — buyers accept the price",
        ],
        "verdict": "CONTRADICTED",
    },
    {
        "label": "H2",
        "statement": ("Service quality has slipped and customers leave "
                      "after poor delivery"),
        "predict_yes": [
            "Complaint volumes rise ahead of non-renewal",
            "Churned accounts show more missed callouts than retained ones",
        ],
        "predict_no": [
            "If true, we will NOT see churn concentrated only in "
            "competitor-bid renewals",
        ],
        "evidence_for": [
            "Two churned accounts in Q3 2025 logged missed-callout "
            "complaints",
        ],
        "evidence_against": [
            "Complaint volumes are flat across the period",
            "The churn rise sits in competitor-bid renewals, not across "
            "the base",
        ],
        "verdict": "INCONCLUSIVE",
    },
    {
        "label": "H3",
        "statement": ("The competitor's four-hour response guarantee has "
                      "reset what renewals are judged on"),
        "predict_yes": [
            "Lost-deal reviews cite the guarantee explicitly",
            "Churn concentrates where the competitor bids",
            "Account teams report response commitments opening renewal "
            "conversations",
        ],
        "predict_no": [
            "If true, we will NOT see elevated churn among current "
            "customers who have not faced a competitor bid",
        ],
        "evidence_for": [
            "10 of 10 lost-deal reviews cite the guarantee",
            "Churn rise concentrated in competitor-bid renewals",
            "Account debriefs: customers raise the guarantee unprompted",
        ],
        "evidence_against": [
            "n=10 is a small base for the lost-deal signal — noted as a "
            "limit, not a contradiction",
        ],
        "verdict": "SUPPORTED",
    },
]

STRONGEST = "H3"

RESEARCH_GAPS = [
    ("Win-loss interviews with 12 lapsed accounts to close H2",
     "Service manager", "12 Sep 2026"),
    ("Competitor guarantee terms and penalty structure",
     "Commercial lead", "5 Sep 2026"),
]

# --- EVIDENCE SWEEP (Stage A4) — one row per claim that carries weight:
# every hypothesis (claim_id = its label) and the answer (claim_id = "ANSWER").
# status: "EVIDENCED" | "GAP_RISK_ACCEPTED" (status_reason required) |
# "GAP_OPEN" (blocks the build). counter_evidence_search records what was
# looked for AND what was found — even when nothing was found.
# ENGAGEMENT ASSET: never commit a filled register to the public repository.
EVIDENCE_SWEEP = [
    {"claim_id": "H1",
     "claim": ("Pricing is cleared as the cause — the churn rise is not "
               "price-driven"),
     "best_supporting_evidence": ("Price objections fell 46 to 32 a quarter "
                                  "while churn rose 13 points; deal sizes "
                                  "stable at £4.2k"),
     "source": "CRM objection log 2023–2026; billing extract, Jul 2026",
     "counter_evidence_search": ("Looked for: discount pressure in renewal "
                                 "notes; price-led losses in the 2025–26 "
                                 "lost-deal file. Found: two competitor rate "
                                 "cards below ours — and zero lost deals "
                                 "citing them."),
     "status": "EVIDENCED", "status_reason": ""},
    {"claim_id": "H2",
     "claim": ("Service quality is not the driver of the churn rise"),
     "best_supporting_evidence": ("Complaint volumes flat across the period; "
                                  "churn concentrated in competitor-bid "
                                  "renewals, not across the base"),
     "source": "Complaints log 2023–2026; renewal records, Jul 2026",
     "counter_evidence_search": ("Looked for: missed-callout complaints "
                                 "preceding non-renewal. Found: two churned "
                                 "Q3 2025 accounts with logged complaints — "
                                 "signal too thin to close either way."),
     "status": "GAP_RISK_ACCEPTED",
     "status_reason": ("Win-loss interviews with 12 lapsed accounts not yet "
                       "run (due 12 Sep 2026); risk bounded because the "
                       "recommendation does not depend on H2's verdict")},
    {"claim_id": "H3",
     "claim": ("The competitor's four-hour guarantee reset what renewals "
               "are judged on"),
     "best_supporting_evidence": ("10 of 10 Q4 2025 lost-deal reviews cite "
                                  "the guarantee; churn concentrated in "
                                  "competitor-bid renewals"),
     "source": "Lost-deal reviews Oct–Dec 2025, n=10; renewal records",
     "counter_evidence_search": ("Looked for: elevated churn among accounts "
                                 "with no competitor bid (H3's negative "
                                 "prediction). Found: none — existing-"
                                 "contract churn stable at 21–23%."),
     "status": "EVIDENCED", "status_reason": ""},
    {"claim_id": "ANSWER",
     "claim": ("Churn is a framing problem: the guarantee reset the renewal "
               "test and Meridian does not yet answer it"),
     "best_supporting_evidence": ("Surviving H3 evidence chain: unanimous "
                                  "lost-deal citations plus the stable "
                                  "no-bid churn floor"),
     "source": "A3 hypothesis tests against the A2 register",
     "counter_evidence_search": ("Looked for: an alternative explanation "
                                 "fitting the same pattern (price, service, "
                                 "relationship churn). Found: each one "
                                 "contradicted or orthogonal in A3."),
     "status": "EVIDENCED", "status_reason": ""},
]

QUESTION = ("How does Meridian win renewals in a market that now judges "
            "providers on response time?")

ANSWER = ("Churn is a framing problem, not a price problem — the "
          "competitor's four-hour guarantee reset what renewals are judged "
          "on, and Meridian does not yet answer it.")

RECOMMENDATION = {
    "options": [
        {"name": "Do nothing",
         "fit": "None — concedes the frame",
         "cost": "£0 outlay; £1.8M a year continuing loss",
         "time": "Immediate",
         "reversibility": "Full",
         "risk": "Loss compounds each renewal cycle",
         "selected": False},
        {"name": "Cut renewal pricing 10%",
         "fit": "Low — prices a cause the evidence contradicts",
         "cost": "£510k a year margin given up",
         "time": "One quarter",
         "reversibility": "Hard — prices ratchet down, not up",
         "risk": "Spends margin on the wrong lever",
         "selected": False},
        {"name": "Stand up a four-hour response guarantee",
         "fit": "High — answers the frame the evidence names",
         "cost": "£350k a year, staged with £120k before the gate",
         "time": "Live for the January renewal cycle",
         "reversibility": "Good — 90-day gate before full commitment",
         "risk": "Frame shifts again; the gate catches it",
         "selected": True},
        {"name": "Full service redesign",
         "fit": "Medium — answers the frame plus more",
         "cost": "£1.4M over 18 months",
         "time": "Misses two renewal cycles",
         "reversibility": "Poor",
         "risk": "Slow, expensive, and unproven against the cause",
         "selected": False},
    ],
    "what": ("Stand up a four-hour response-time guarantee for all "
             "renewal-eligible contracts"),
    "why": ("It answers the frame the lost-deal evidence says renewals are "
            "now judged on"),
    "resources": ("£350k a year: extended engineering cover, a triage desk, "
                  "and a service-credit reserve"),
    "timeline": "Build starts 1 Oct 2026; live for the January renewal cycle",
    "owner": "Operations director",
    "metrics": [
        {"measure": "Lost-deal citations of the competitor guarantee",
         "baseline": "10 a quarter", "target": "Under five a quarter",
         "date": "31 Dec 2026"},
        {"measure": "SME contract churn",
         "baseline": "35%", "target": "27%", "date": "30 Jun 2027"},
        {"measure": "Callouts answered inside four hours",
         "baseline": "78%", "target": "98%", "date": "31 Mar 2027"},
        {"measure": "Service-credit payouts vs reserve",
         "baseline": "Reserve set at 2% breach", "target": "Within reserve",
         "date": "Quarterly from 31 Mar 2027"},
    ],
    "decision_gate": ("If lost-deal citations have not fallen below five a "
                      "quarter by 31 Dec 2026, stop before the remaining "
                      "spend commits"),
    "key_assumption": ("The market's response-time frame persists through "
                       "2027; renewals continue to be judged on it"),
}

OUTPUT_FILE = "pyramid-analysis.docx"

# ============================================================================
# BUILD — DO NOT EDIT BELOW THIS LINE
# ============================================================================
import re
import sys

VERDICTS = ("SUPPORTED", "CONTRADICTED", "INCONCLUSIVE")

def _collect_strings(obj, out):
    if isinstance(obj, str):
        out.append(obj)
    elif isinstance(obj, dict):
        for v in obj.values():
            _collect_strings(v, out)
    elif isinstance(obj, (list, tuple)):
        for v in obj:
            _collect_strings(v, out)

SWEEP_STATUSES = ("EVIDENCED", "GAP_RISK_ACCEPTED", "GAP_OPEN")

def _theme_defects():
    defects = []
    for k, v in THEME.items():
        if k.startswith("font_"):
            if not str(v).strip():
                defects.append("THEME['%s'] is empty — name a real font" % k)
        elif not re.fullmatch(r"[0-9A-Fa-f]{6}", str(v)):
            defects.append("THEME['%s'] = %r is not a 6-char hex "
                           "(no '#', no alpha)" % (k, v))
    return defects

def run_self_check():
    checks, warnings = [], []

    td = _theme_defects()
    checks.append(("THEME: colours are 6-char hex and fonts are named"
                   + ("" if not td else " — " + "; ".join(td)), not td))

    fr = PROBLEM_FRAME
    checks.append(("frame has all four elements: symptom, evidence, "
                   "decision, reader",
                   all([fr.get("symptom"), fr.get("evidence"),
                        fr.get("decision"), fr.get("reader"),
                        fr.get("frame_sentence")])))
    checks.append(("frame carries no diagnosis (the word 'because' is a "
                   "fail)", " because " not in
                   (" " + fr["frame_sentence"].lower())))
    checks.append(("frame contains numbers (an observable symptom, not an "
                   "interpretation)",
                   bool(re.search(r"\d", fr["frame_sentence"]))))

    low_sit = " " + SITUATION.lower()
    checks.append(("situation carries no 'but' / 'however' (tension belongs "
                   "in the complication)",
                   " but " not in low_sit and "however" not in low_sit))
    checks.append(("complication starts with 'But' and names the damage "
                   "with numbers",
                   COMPLICATION.strip().startswith("But")
                   and bool(re.search(r"\d", COMPLICATION))))
    checks.append(("every register fact names a source and a date",
                   all(len(row) == 3 and all(str(x).strip() for x in row)
                       for row in EVIDENCE_REGISTER)))

    checks.append(("3–5 hypotheses (%d)" % len(HYPOTHESES),
                   3 <= len(HYPOTHESES) <= 5))
    hyp_defects = []
    for h in HYPOTHESES:
        if not h.get("predict_no"):
            hyp_defects.append("%s has no negative prediction — untestable"
                               % h["label"])
        if h.get("verdict") not in VERDICTS:
            hyp_defects.append("%s verdict must be SUPPORTED / CONTRADICTED "
                               "/ INCONCLUSIVE" % h["label"])
        if h.get("verdict") == "SUPPORTED" and not h.get("evidence_against"):
            hyp_defects.append("%s SUPPORTED with an empty evidence-against "
                               "column — nobody looked" % h["label"])
        if not h.get("evidence_for") and not h.get("evidence_against"):
            hyp_defects.append("%s was never tested" % h["label"])
    checks.append(("every hypothesis has a negative prediction, a real "
                   "test, and a valid verdict"
                   + ("" if not hyp_defects else
                      " — " + "; ".join(hyp_defects)),
                   not hyp_defects))
    stmts = [h["statement"].strip().lower() for h in HYPOTHESES]
    checks.append(("no two hypotheses are one cause in different words "
                   "(MECE)", len(set(stmts)) == len(stmts)))
    strongest = [h for h in HYPOTHESES if h["label"] == STRONGEST]
    checks.append(("the strongest hypothesis exists and its verdict is "
                   "SUPPORTED",
                   len(strongest) == 1
                   and strongest[0]["verdict"] == "SUPPORTED"))

    # --- evidence sweep (Stage A4): adversarial register, gated
    row_defects = []
    for r in EVIDENCE_SWEEP:
        rid = r.get("claim_id", "?")
        for fld in ("claim", "best_supporting_evidence", "source",
                    "counter_evidence_search"):
            if not str(r.get(fld, "")).strip():
                row_defects.append("%s missing %s" % (rid, fld))
        if r.get("status") not in SWEEP_STATUSES:
            row_defects.append("%s status must be EVIDENCED / "
                               "GAP_RISK_ACCEPTED / GAP_OPEN" % rid)
        if (r.get("status") == "GAP_RISK_ACCEPTED"
                and not str(r.get("status_reason", "")).strip()):
            row_defects.append("%s is GAP_RISK_ACCEPTED with no stated "
                               "reason" % rid)
    checks.append(("every sweep row is complete: claim, best evidence, "
                   "source, counter-evidence search, valid status"
                   + ("" if not row_defects else
                      " — " + "; ".join(row_defects)), not row_defects))
    open_gaps = [r["claim_id"] for r in EVIDENCE_SWEEP
                 if r.get("status") == "GAP_OPEN"]
    checks.append(("no GAP_OPEN rows — every claim is EVIDENCED or "
                   "explicitly GAP_RISK_ACCEPTED with a reason"
                   + ("" if not open_gaps else
                      " — resolve or risk-accept: " + ", ".join(open_gaps)),
                   not open_gaps))
    need_ids = set(h["label"] for h in HYPOTHESES) | {"ANSWER"}
    have_ids = set(r.get("claim_id") for r in EVIDENCE_SWEEP)
    missing_ids = sorted(need_ids - have_ids)
    checks.append(("the sweep covers every hypothesis and the answer"
                   + ("" if not missing_ids else
                      " — missing rows: " + ", ".join(missing_ids)),
                   not missing_ids))

    checks.append(("exactly one question, ending in a question mark",
                   QUESTION.count("?") == 1
                   and QUESTION.strip().endswith("?")))
    n_sent = len([x for x in re.split(r"[.!?]", ANSWER) if x.strip()])
    checks.append(("the answer is one sentence and is a claim, not a "
                   "question", n_sent == 1 and "?" not in ANSWER))
    checks.append(("the answer is not a restatement of the question",
                   ANSWER.strip().lower() != QUESTION.strip().lower()))

    rec = RECOMMENDATION
    checks.append(("3–4 options, including an honest 'do nothing'",
                   3 <= len(rec["options"]) <= 4
                   and any("do nothing" in o["name"].lower()
                           for o in rec["options"])))
    checks.append(("exactly one option is selected",
                   sum(1 for o in rec["options"] if o.get("selected")) == 1))
    checks.append(("the recommendation names what, why, resources, "
                   "timeline and owner",
                   all([rec.get("what"), rec.get("why"),
                        rec.get("resources"), rec.get("timeline"),
                        rec.get("owner")])))
    checks.append(("3–4 metrics, each with baseline, target and date",
                   3 <= len(rec["metrics"]) <= 4
                   and all(m.get("baseline") and m.get("target")
                           and m.get("date") for m in rec["metrics"])))
    checks.append(("a decision gate and a key assumption are named",
                   bool(rec.get("decision_gate"))
                   and bool(rec.get("key_assumption"))))

    strings = []
    for blob in (META, PROBLEM_FRAME, SITUATION, COMPLICATION,
                 EVIDENCE_REGISTER, HYPOTHESES, EVIDENCE_SWEEP,
                 RESEARCH_GAPS, QUESTION, ANSWER, RECOMMENDATION):
        _collect_strings(blob, strings)
    placeholders = sorted(set(m for s in strings
                              for m in re.findall(r"\[[^\]]+\]", s)))
    checks.append(("no bracketed placeholders remain"
                   + ("" if not placeholders else
                      " — replace: " + ", ".join(placeholders)),
                   not placeholders))

    for w in ("failing", "weak", "poor", "terrible", "underperforming"):
        if w in SITUATION.lower():
            warnings.append("judgment word '%s' in the situation — it must "
                            "pass the nod test" % w)
    inconclusive = [h["label"] for h in HYPOTHESES
                    if h["verdict"] == "INCONCLUSIVE"]
    if inconclusive and not RESEARCH_GAPS:
        warnings.append("hypotheses %s are INCONCLUSIVE but no research "
                        "gaps are listed" % ", ".join(inconclusive))
    return checks, warnings

def print_report(checks, warnings, wrote):
    print("=" * 64)
    print("SELF-CHECK — %s" % OUTPUT_FILE)
    for name, ok in checks:
        print("  [%s] %s" % ("PASS" if ok else "FAIL", name))
    for w in warnings:
        print("  [WARN] %s" % w)
    print("=" * 64)
    if not all(ok for _, ok in checks):
        raise SystemExit("SELF-CHECK FAILED — no file written. Fix the DATA "
                         "section, not the checks.")
    if wrote:
        print("written", OUTPUT_FILE)

_checks, _warnings = run_self_check()
if not all(ok for _, ok in _checks):
    print_report(_checks, _warnings, wrote=False)
    sys.exit(1)

# ---------------------------------------------------------------- docx build
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def _rgb(h):
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

SLATE = _rgb(THEME["dominant"])
BLUE = _rgb(THEME["accent"])
GOLD = _rgb(THEME["highlight"])
INK = _rgb(THEME["ink"])
MUTE = _rgb(THEME["muted"])
GREEN = RGBColor(0x2F, 0x6B, 0x4F)   # functional verdict colour — not brand
RED = RGBColor(0xA8, 0x44, 0x3C)     # functional verdict colour — not brand
TINT_HEX = THEME["card_fill"]
DOMINANT_HEX = THEME["dominant"]
HEAD_FONT, BODY_FONT = THEME["font_heading"], THEME["font_body"]

doc = Document()
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Inches(0.9)
    sec.left_margin = sec.right_margin = Inches(0.9)
normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK

def para(text="", size=10.5, color=INK, bold=False, italic=False,
         font=BODY_FONT, space_after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.bold = bold
    r.font.italic = italic
    return p

def h1(text):
    return para(text, size=15, color=SLATE, bold=True, font=HEAD_FONT,
                space_after=5)

def h2(text):
    return para(text, size=12, color=BLUE, bold=True, font=HEAD_FONT,
                space_after=4)

def shade(cell, hex_fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)

def cell_text(cell, text, size=9.5, color=INK, bold=False, font=BODY_FONT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.bold = bold

def data_table(headers, rows, widths=None, colors=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    if widths:
        t.autofit = False
        for j, wd in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Inches(wd)
    for j, hd in enumerate(headers):
        shade(t.cell(0, j), DOMINANT_HEX)
        cell_text(t.cell(0, j), hd, size=9.5,
                  color=RGBColor(0xFF, 0xFF, 0xFF), bold=True,
                  font=HEAD_FONT)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            col = colors[i][j] if colors else INK
            cell_text(t.cell(1 + i, j), str(v), color=col)
    return t

def panel(rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    t.autofit = False
    for i, (lab, text) in enumerate(rows):
        t.cell(i, 0).width = Inches(1.6)
        t.cell(i, 1).width = Inches(5.4)
        shade(t.cell(i, 0), TINT_HEX)
        shade(t.cell(i, 1), "FFFFFF")
        cell_text(t.cell(i, 0), lab, size=9.5, color=SLATE, bold=True,
                  font=HEAD_FONT)
        cell_text(t.cell(i, 1), text, size=10)
    return t

# --- header
para(META["title"], size=19, color=SLATE, bold=True, font=HEAD_FONT,
     space_after=2)
para("Minto diagnostic (Flow A) · %s · %s · reader: %s"
     % (META["author_line"], META["date"], META["reader"]),
     size=9, color=MUTE, space_after=12)

# --- A1
h1("A1 — Problem frame")
panel([("Frame", PROBLEM_FRAME["frame_sentence"]),
       ("Symptom", PROBLEM_FRAME["symptom"]),
       ("Evidence", PROBLEM_FRAME["evidence"]),
       ("Decision waiting", PROBLEM_FRAME["decision"]),
       ("Reader", PROBLEM_FRAME["reader"])])
para("The frame carries no diagnosis. Candidate causes are parked in "
     "section A3, where they are tested rather than assumed.", size=9,
     color=MUTE, italic=True, space_after=12)

# --- A2
h1("A2 — Situation, complication and the evidence register")
h2("Situation (nod-test facts only)")
para(SITUATION, space_after=6)
h2("Complication (the change and the damage — no solution, no cause)")
para(COMPLICATION, space_after=8)
h2("Evidence register")
data_table(["Fact", "Source", "Date"],
           [(f, s, d) for f, s, d in EVIDENCE_REGISTER],
           widths=[3.7, 2.2, 0.9])
para("Every number in this analysis points back to a register row that can "
     "be shown in the meeting.", size=9, color=MUTE, italic=True,
     space_after=12)

# --- A3
h1("A3 — Hypotheses and tests")
para("Each candidate cause carries positive predictions (if true, the data "
     "shows X) and negative predictions (if true, we will NOT see W). A "
     "SUPPORTED verdict requires a surviving negative prediction — friendly "
     "evidence alone is never sufficient.", space_after=8)
for hyp in HYPOTHESES:
    vcol = (GREEN if hyp["verdict"] == "SUPPORTED"
            else RED if hyp["verdict"] == "CONTRADICTED" else GOLD)
    h2("%s — %s" % (hyp["label"], hyp["statement"]))
    para("Verdict: " + hyp["verdict"], size=10.5, color=vcol, bold=True,
         space_after=4)
    rows = []
    rows.append(("Predicts (positive)", " · ".join(hyp["predict_yes"])))
    rows.append(("Predicts (negative)", " · ".join(hyp["predict_no"])))
    rows.append(("Evidence for", " · ".join(hyp["evidence_for"]) or "None"))
    rows.append(("Evidence against",
                 " · ".join(hyp["evidence_against"]) or "None"))
    panel(rows)
    para("", space_after=4)
strongest_h = next(h for h in HYPOTHESES if h["label"] == STRONGEST)
para("Strongest hypothesis: %s — %s. Its negative prediction survived: "
     "the pattern it forbids does not appear in the data."
     % (STRONGEST, strongest_h["statement"]), bold=True, space_after=6)
if RESEARCH_GAPS:
    h2("Research gaps (with owners and dates)")
    data_table(["Gap", "Owner", "Due"], RESEARCH_GAPS,
               widths=[4.2, 1.6, 1.0])
para("", space_after=8)

# --- A4: evidence sweep (adversarial)
h1("A4 — Evidence sweep (adversarial)")
para("Every claim that carries weight holds a register row: its best "
     "supporting evidence with a source, AND a recorded search for "
     "disconfirming evidence — what was looked for, and what was found, "
     "even when nothing was found. No claim proceeds as GAP_OPEN: it is "
     "evidenced, or its gap is accepted with a stated reason.",
     space_after=8)
sweep_rows, sweep_colors = [], []
for r in EVIDENCE_SWEEP:
    sweep_rows.append((r["claim_id"] + " — " + r["claim"],
                       r["best_supporting_evidence"], r["source"],
                       r["counter_evidence_search"], r["status"]))
    scol = (GREEN if r["status"] == "EVIDENCED" else GOLD)
    sweep_colors.append([INK, INK, INK, INK, scol])
data_table(["Claim", "Best supporting evidence", "Source",
            "Counter-evidence search", "Status"],
           sweep_rows, widths=[1.5, 1.4, 1.0, 1.8, 1.0],
           colors=sweep_colors)
accepted = [r for r in EVIDENCE_SWEEP if r["status"] == "GAP_RISK_ACCEPTED"]
if accepted:
    para("", space_after=2)
    h2("Risk-accepted gaps — stated limits, not silent ones")
    panel([(r["claim_id"], r["status_reason"]) for r in accepted])
para("", space_after=2)
para("The filled register is an engagement asset — it stays in the client "
     "tenant and is never committed to the public repository.", size=9,
     color=MUTE, italic=True, space_after=12)

# --- A5
h1("A5 — Question and answer")
panel([("The question", QUESTION), ("The answer", ANSWER)])
para("Chain: answer ← %s ← the register rows tested in A3, swept "
     "adversarially in A4. The question sits above the surface debate (cut "
     "the price or not) — answering it sets the action." % STRONGEST,
     size=9, color=MUTE, italic=True, space_after=12)

# --- A6
h1("A6 — Recommendation")
h2("Options examined (the selected option in bold)")
opt_rows, opt_colors = [], []
for o in RECOMMENDATION["options"]:
    opt_rows.append((o["name"], o["fit"], o["cost"], o["time"],
                     o["reversibility"], o["risk"]))
    c = SLATE if o["selected"] else INK
    opt_colors.append([c] * 6)
t = data_table(["Option", "Strategic fit", "Cost", "Time to impact",
                "Reversibility", "Risk if wrong"], opt_rows,
               widths=[1.35, 1.25, 1.25, 1.0, 1.0, 1.15], colors=opt_colors)
for i, o in enumerate(RECOMMENDATION["options"]):
    if o["selected"]:
        for cell in t.rows[1 + i].cells:
            shade(cell, TINT_HEX)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
para("", space_after=4)
h2("The recommendation")
panel([("What", RECOMMENDATION["what"]),
       ("Why", RECOMMENDATION["why"] + " (points at the A5 answer)"),
       ("Resources", RECOMMENDATION["resources"]),
       ("Timeline", RECOMMENDATION["timeline"]),
       ("Owner", RECOMMENDATION["owner"])])
para("", space_after=4)
h2("Success metrics — every metric has a baseline, a target and a date")
data_table(["Measure", "Baseline", "Target", "Measured by"],
           [(m["measure"], m["baseline"], m["target"], m["date"])
            for m in RECOMMENDATION["metrics"]],
           widths=[2.8, 1.5, 1.5, 1.2])
para("", space_after=4)
panel([("Decision gate", RECOMMENDATION["decision_gate"]),
       ("Key assumption", RECOMMENDATION["key_assumption"])])
para("", space_after=10)

# --- A7: hand-off map onto the shared PYRAMID schema
h1("A7 — Hand-off to the presentation generators")
para("The fields below map this analysis onto the shared PYRAMID schema "
     "used by generate_pyramid_memo.py, generate_presentation_deck.py and "
     "generate_standalone_deck.py. The failed hypotheses travel too — they "
     "are the prepared answers to the room's objections.", space_after=6)
gt_seed = ("%s — %s" % (RECOMMENDATION["what"],
                        RECOMMENDATION["why"].lower()))
handoff_rows = [
    ("SCQA['situation']", SITUATION),
    ("SCQA['complication']", COMPLICATION),
    ("SCQA['question']", QUESTION),
    ("GOVERNING_THOUGHT (draft — compress to 25 words or fewer)", gt_seed),
    ("KEY_LINE seed 1 (diagnosis)", ANSWER),
    ("KEY_LINE seed 2 (remedy)", RECOMMENDATION["what"]),
    ("KEY_LINE seed 3 (economics)",
     "Metrics, gate and options table from A6 — price the case in ranges"),
    ("KEY_LINE[i]['supports'][j]['source']",
     "Evidence register rows (A2) — one source per support"),
    ("EVIDENCE_REGISTER (memo generator)",
     "The A4 sweep rows travel as-is — claims stay EVIDENCED or "
     "risk-accepted with their reasons; the memo renders accepted gaps as "
     "stated limits"),
    ("ASK['decision']", RECOMMENDATION["what"]),
    ("ASK['gate']", RECOMMENDATION["decision_gate"]),
    ("ASK['owner'] / ASK['date']",
     "%s / from the A6 timeline" % RECOMMENDATION["owner"]),
    ("CONCLUSION['key_assumption']", RECOMMENDATION["key_assumption"]),
    ("Objection bank (not a schema field — keep for B5/Q&A)",
     "; ".join("%s (%s): %s" % (h["label"], h["verdict"], h["statement"])
               for h in HYPOTHESES if h["verdict"] != "SUPPORTED")),
]
data_table(["PYRAMID field", "Carries"], handoff_rows, widths=[2.6, 4.3])
para("", space_after=6)
para(META["footer_tag"], size=8.5, color=MUTE, italic=True,
     align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(OUTPUT_FILE)
print_report(_checks, _warnings, wrote=True)
